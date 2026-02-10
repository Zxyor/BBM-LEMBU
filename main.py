import streamlit as st
import mysql.connector
import pandas as pd
import io
import datetime
import math
import re
from dateutil.relativedelta import relativedelta

# --- SETUP MATPLOTLIB ---
import matplotlib
matplotlib.use('Agg') 
from matplotlib.figure import Figure
from matplotlib.backends.backend_agg import FigureCanvasAgg
import matplotlib.ticker as ticker
import matplotlib.pyplot as plt

# --- LIBRARY REPORTING ---
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, portrait
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.units import cm, mm

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.drawing.image import Image as XLImage

# --- GLOBAL COLORS ---
COLOR_HEADER_BLUE = colors.HexColor("#2F5496")
COLOR_TOTAL_YELLOW = colors.HexColor("#FFD966")
COLOR_ROW_EVEN = colors.HexColor("#F2F2F2")
COLOR_ROW_ODD = colors.white
COLOR_BORDER = colors.HexColor("#000000") 

# --- KONFIGURASI HALAMAN ---
st.set_page_config(page_title="Sistem BBM Proyek LEMBU", layout="wide")

# --- KONEKSI DATABASE ---
def init_connection():
    return mysql.connector.connect(
        host=st.secrets["db"]["host"],
        user=st.secrets["db"]["user"],
        password=st.secrets["db"]["password"],
        database=st.secrets["db"]["database"],
        port=st.secrets["db"]["port"]
    )

# --- HELPER FUNCTIONS ---
def get_bulan_indonesia(bulan_int):
    nama_bulan = ["", "JANUARI", "FEBRUARI", "MARET", "APRIL", "MEI", "JUNI", 
                  "JULI", "AGUSTUS", "SEPTEMBER", "OKTOBER", "NOVEMBER", "DESEMBER"]
    return nama_bulan[bulan_int]

def get_hari_indonesia(tanggal):
    try:
        if pd.isnull(tanggal): return "-"
        kamus = {'Monday': 'Senin', 'Tuesday': 'Selasa', 'Wednesday': 'Rabu', 
                 'Thursday': 'Kamis', 'Friday': 'Jumat', 'Saturday': 'Sabtu', 'Sunday': 'Minggu'}
        return kamus[tanggal.strftime('%A')]
    except: return "-"

def cek_kategori(nama_alat):
    nama = str(nama_alat).upper()
    kata_kunci_mobil = ["TRUCK", "MOBIL", "TRITON", "DT", "FAW", "SANNY", "R6", "R10", "PICK UP", "HILUX", "STRADA", "GRAND MAX"]
    if any(k in nama for k in kata_kunci_mobil):
        return "MOBIL_TRUCK"
    return "ALAT_BERAT"

def set_cell_bg(cell, color_hex):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def segregate_data(df, excluded_list):
    if df.empty:
        empty_df = pd.DataFrame(columns=['nama_alat', 'no_unit', 'jumlah_liter', 'kategori'])
        return empty_df, empty_df, empty_df
    
    df = df.copy()
    df['full_name'] = df['nama_alat'].astype(str) + " " + df['no_unit'].astype(str)
    
    df_lainnya = df[df['full_name'].isin(excluded_list)].copy()
    df_main = df[~df['full_name'].isin(excluded_list)].copy()
    
    df_alat = df_main[df_main['kategori'] == 'ALAT_BERAT']
    df_truck = df_main[df_main['kategori'] == 'MOBIL_TRUCK']
    
    return df_alat, df_truck, df_lainnya

def filter_non_consumption(df):
    if df.empty: return df
    mask_donor = df['jumlah_liter'] < 0
    mask_recv = df['keterangan'].str.contains('Transfer|Pinjam', case=False, na=False)
    return df[~(mask_donor | mask_recv)]

def hitung_stok_awal_periode(conn, lokasi_id, start_date):
    cursor = conn.cursor()
    cursor.execute("SELECT stok_awal FROM lokasi_proyek WHERE id = %s", (lokasi_id,))
    res = cursor.fetchone()
    modal_awal = float(res[0]) if res and res[0] is not None else 0.0
    cursor.execute("SELECT COALESCE(SUM(jumlah_liter), 0) FROM bbm_masuk WHERE lokasi_id = %s AND tanggal < %s", (lokasi_id, start_date))
    res_m = cursor.fetchone()
    masuk_prev = float(res_m[0])
    cursor.execute("SELECT COALESCE(SUM(jumlah_liter), 0) FROM bbm_keluar WHERE lokasi_id = %s AND tanggal < %s", (lokasi_id, start_date))
    res_k = cursor.fetchone()
    keluar_prev = float(res_k[0])
    return modal_awal + masuk_prev - keluar_prev

def split_date_range_by_month(start_date, end_date):
    result = []
    current = start_date.replace(day=1)
    while current <= end_date:
        month_start = max(start_date, current)
        next_month = current + relativedelta(months=1)
        month_end = min(end_date, next_month - datetime.timedelta(days=1))
        if month_start <= month_end:
            result.append((month_start, month_end))
        current = next_month
    return result

def safe_text(text, max_chars=25):
    s = str(text) if text else "-"
    if len(s) > max_chars:
        return s[:max_chars] + "..."
    return s

# --- CHART GENERATOR ---
def generate_chart_for_report(df_alat, df_truck, width_inch=6, height_inch=3):
    try:
        active_charts = []
        if not df_alat.empty: active_charts.append(("PEMAKAIAN ALAT BERAT", df_alat, '#F4B084'))
        if not df_truck.empty: active_charts.append(("PEMAKAIAN MOBIL & TRUCK", df_truck, '#9BC2E6'))
        
        num_charts = len(active_charts)
        if num_charts == 0: return None
        
        # Adjust height
        total_height = height_inch * num_charts
        fig = Figure(figsize=(width_inch, total_height), dpi=150)
        canvas = FigureCanvasAgg(fig)
        axs = fig.subplots(num_charts, 1)
        if num_charts == 1: axs = [axs]
        
        for i, (title, df, color) in enumerate(active_charts):
            ax = axs[i]
            rekap = df.groupby(['nama_alat', 'no_unit'])['jumlah_liter'].sum().reset_index()
            data = rekap.sort_values('jumlah_liter', ascending=True)
            labels = data.apply(lambda x: f"{x['nama_alat']} {x['no_unit']}", axis=1)
            bars = ax.barh(labels, data['jumlah_liter'], color=color, edgecolor='#555555', height=0.7)
            ax.set_title(title, fontsize=10, fontweight='bold')
            ax.tick_params(labelsize=8)
            ax.xaxis.set_major_formatter(ticker.FuncFormatter(lambda x, p: format(int(x), ',')))
            for bar in bars:
                width = bar.get_width()
                ax.text(width, bar.get_y() + bar.get_height()/2, f' {width:,.0f}', va='center', fontsize=8)

        fig.tight_layout()
        buf = io.BytesIO()
        canvas.print_png(buf)
        buf.seek(0)
        return buf
    except: return None

def generate_monthly_chart(df_monthly):
    try:
        if df_monthly.empty: return None
        masuk_vals = pd.to_numeric(df_monthly['masuk'], errors='coerce').fillna(0)
        keluar_vals = pd.to_numeric(df_monthly['keluar'], errors='coerce').fillna(0)
        labels = df_monthly['bulan_nama'].astype(str).tolist()

        fig = Figure(figsize=(8, 4), dpi=100)
        canvas = FigureCanvasAgg(fig)
        ax = fig.add_subplot(111)
        x = range(len(labels))
        width = 0.35
        ax.bar([i - width/2 for i in x], masuk_vals, width, label='Masuk', color='#90EE90', edgecolor='black')
        ax.bar([i + width/2 for i in x], keluar_vals, width, label='Keluar', color='#F08080', edgecolor='black')
        ax.set_title('GRAFIK MASUK & PENGGUNAAN BBM PER BULAN', fontsize=10, fontweight='bold')
        ax.set_xticks(x)
        ax.set_xticklabels(labels, fontsize=8)
        ax.legend(fontsize=8)
        ax.yaxis.set_major_formatter(ticker.FuncFormatter(lambda x, p: format(int(x), ',')))
        ax.tick_params(labelsize=8)
        fig.tight_layout()
        buf = io.BytesIO()
        canvas.print_png(buf)
        buf.seek(0)
        return buf
    except: return None

# ==========================================
# 1. EXPORT STANDARD (A4 PAGE BREAK)
# ==========================================

def generate_pdf_portrait(conn, lokasi_id, nama_lokasi, start_date_global, end_date_global, excluded_list):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=portrait(A4), rightMargin=15, leftMargin=15, topMargin=20, bottomMargin=20)
    elements = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(name='ExcelTitle', parent=styles['Heading1'], alignment=TA_CENTER, fontSize=14, fontName='Helvetica-Bold', spaceAfter=2, textColor=colors.HexColor("#2F5496"))
    periode_style = ParagraphStyle(name='ExcelPeriode', parent=styles['Normal'], alignment=TA_CENTER, fontSize=11, spaceAfter=15, textColor=colors.black)
    cell_style = ParagraphStyle(name='CellText', parent=styles['Normal'], fontSize=7, leading=8, fontName='Helvetica')
    header_style = ParagraphStyle(name='HeaderTxt', parent=styles['Normal'], fontSize=7, leading=8, fontName='Helvetica-Bold', textColor=colors.white, alignment=TA_CENTER)
    header_black_style = ParagraphStyle(name='HeaderTxtBlk', parent=styles['Normal'], fontSize=7, leading=8, fontName='Helvetica-Bold', textColor=colors.black, alignment=TA_CENTER)
    section_title_style = ParagraphStyle(name='SectionTitle', parent=styles['Normal'], fontSize=8, leading=9, fontName='Helvetica-Bold', textColor=colors.HexColor("#2F5496"))
    
    date_ranges = split_date_range_by_month(start_date_global, end_date_global)
    for idx, (start_date, end_date) in enumerate(date_ranges):
        if idx > 0: elements.append(PageBreak())
        stok_awal = hitung_stok_awal_periode(conn, lokasi_id, start_date)
        df_masuk = pd.read_sql(f"SELECT * FROM bbm_masuk WHERE lokasi_id={lokasi_id} AND tanggal BETWEEN '{start_date}' AND '{end_date}' ORDER BY tanggal", conn)
        df_keluar = pd.read_sql(f"SELECT * FROM bbm_keluar WHERE lokasi_id={lokasi_id} AND tanggal BETWEEN '{start_date}' AND '{end_date}' ORDER BY tanggal", conn)
        if not df_keluar.empty and 'kategori' not in df_keluar.columns: df_keluar['kategori'] = df_keluar['nama_alat'].apply(cek_kategori)
        df_keluar_rpt = filter_non_consumption(df_keluar)
        df_alat_g, df_truck_g, df_lain_g = segregate_data(df_keluar_rpt, excluded_list)
        tm = float(df_masuk['jumlah_liter'].sum()) if not df_masuk.empty else 0.0
        tk_real = float(df_keluar['jumlah_liter'].sum()) if not df_keluar.empty else 0.0
        tk_rpt = float(df_keluar_rpt['jumlah_liter'].sum()) if not df_keluar_rpt.empty else 0.0
        sisa_akhir = stok_awal + tm - tk_real

        elements.append(Paragraph(f"LAPORAN BBM: {nama_lokasi}", title_style))
        elements.append(Paragraph(f"PERIODE {get_bulan_indonesia(start_date.month)} {start_date.year}", periode_style))
        
        left_queue = []; left_queue.append({'type': 'title_section', 'val': 'PENGGUNAAN BBM (KELUAR)'}); left_queue.append({'type': 'header_col'}) 
        if not df_keluar_rpt.empty:
            for i, r in df_keluar_rpt.iterrows(): left_queue.append({'type': 'row', 'data': [i+1, r['tanggal'].strftime('%d/%m'), r['nama_alat'], r['no_unit'], f"{r['jumlah_liter']:.0f}", r['keterangan']], 'date_val': r['tanggal']})
        left_queue.append({'type': 'total_left', 'val': f"{tk_rpt:.0f}"})

        right_queue = []; right_queue.append({'type': 'title_section', 'val': 'BBM MASUK'}); right_queue.append({'type': 'header_masuk'})
        if not df_masuk.empty:
            for i, r in df_masuk.iterrows(): right_queue.append({'type': 'row_masuk', 'data': [i+1, r['tanggal'].strftime('%d/%m'), r['sumber'], r['jenis_bbm'], f"{r['jumlah_liter']:.0f}"]})
        else: right_queue.append({'type': 'row_masuk', 'data': ['-', '-', 'TIDAK ADA DATA', '-', '0']})
        right_queue.append({'type': 'total_masuk', 'val': f"{tm:.0f}"})
        
        right_queue.append({'type': 'title_section', 'val': 'RINCIAN PENGGUNAAN BBM'})
        def add_rekap(df, title, color, text_is_black=False):
            right_queue.append({'type': 'sub_rekap', 'title': title, 'bg': color, 'txt_black': text_is_black})
            if not df.empty:
                grp = df.groupby(['nama_alat', 'no_unit'])['jumlah_liter'].sum().reset_index().sort_values('jumlah_liter', ascending=False)
                for _, r in grp.iterrows(): right_queue.append({'type': 'row_rekap', 'label': f"{r['nama_alat']} {r['no_unit']}", 'val': f"{r['jumlah_liter']:.0f}"})
                right_queue.append({'type': 'total_rekap', 'val': f"{df['jumlah_liter'].sum():.0f}"})
            else: right_queue.append({'type': 'row_rekap', 'label': '-', 'val': '0'}); right_queue.append({'type': 'total_rekap', 'val': '0'})

        add_rekap(df_alat_g, "TOTAL ALAT BERAT", "#F4B084", True); add_rekap(df_truck_g, "TOTAL MOBIL & TRUCK", "#9BC2E6", True)
        if not df_lain_g.empty: add_rekap(df_lain_g, "LAINNYA", "#ED77C4", False)
        
        right_queue.append({'type': 'title_section', 'val': 'RINCIAN SISA STOK BBM'}); right_queue.append({'type': 'header_stok', 'label': 'RINGKASAN STOK'})
        right_queue.append({'type': 'row_stok', 'label': 'SISA BULAN LALU', 'val': f"{stok_awal:.0f}"}); right_queue.append({'type': 'row_stok', 'label': 'TOTAL MASUK', 'val': f"{tm:.0f}"})
        right_queue.append({'type': 'row_stok', 'label': 'TOTAL KELUAR', 'val': f"{tk_real:.0f}"}); right_queue.append({'type': 'total_stok', 'label': 'SISA AKHIR', 'val': f"{sisa_akhir:.0f}"})
        
        img_buf = generate_chart_for_report(df_alat_g, df_truck_g, width_inch=3.5, height_inch=2.5)
        if img_buf: 
            num_charts = (1 if not df_alat_g.empty else 0) + (1 if not df_truck_g.empty else 0)
            right_queue.append({'type': 'chart', 'img': img_buf, 'span': 15 * num_charts})

        ROWS_PER_PAGE = 40; ROW_HEIGHT = 15; l_ptr = 0; r_ptr = 0; right_occupied_until = -1
        last_date_zebra = None; is_zebra_grey = False 
        while True:
            page_data = []; page_style = [('VALIGN', (0,0), (-1,-1), 'MIDDLE')]; row_idx = 0
            while row_idx < ROWS_PER_PAGE:
                if l_ptr >= len(left_queue) and r_ptr >= len(right_queue) and row_idx > right_occupied_until: break
                row_content = [''] * 12 
                if l_ptr < len(left_queue):
                    item = left_queue[l_ptr]; itype = item['type']
                    if itype == 'title_section': row_content[0] = Paragraph(item['val'], section_title_style); page_style.append(('SPAN', (0, row_idx), (5, row_idx)))
                    elif itype == 'header_col':
                        cols = ['NO', 'TGL', 'ALAT', 'UNIT', 'LTR', 'KET']
                        for c, txt in enumerate(cols): row_content[c] = Paragraph(txt, header_style)
                        page_style.append(('BACKGROUND', (0, row_idx), (5, row_idx), COLOR_HEADER_BLUE)); page_style.append(('GRID', (0, row_idx), (5, row_idx), 0.5, COLOR_BORDER))
                    elif itype == 'row':
                        d = item['data']; row_content[0] = d[0]; row_content[1] = d[1]; row_content[2] = Paragraph(safe_text(d[2], 25), cell_style); row_content[3] = Paragraph(safe_text(d[3], 15), cell_style); row_content[4] = d[4]; row_content[5] = Paragraph(safe_text(d[5], 25), cell_style)
                        
                        curr_date = item.get('date_val')
                        if last_date_zebra is not None and curr_date != last_date_zebra: is_zebra_grey = not is_zebra_grey
                        last_date_zebra = curr_date
                        bg = COLOR_ROW_EVEN if is_zebra_grey else COLOR_ROW_ODD
                        
                        page_style.append(('BACKGROUND', (0, row_idx), (5, row_idx), bg)); page_style.append(('GRID', (0, row_idx), (5, row_idx), 0.5, COLOR_BORDER))
                    elif itype == 'total_left':
                        row_content[2] = 'TOTAL PENGGUNAAN'; row_content[4] = item['val']; page_style.append(('SPAN', (2, row_idx), (3, row_idx))); page_style.append(('BACKGROUND', (0, row_idx), (5, row_idx), COLOR_TOTAL_YELLOW)); page_style.append(('FONTNAME', (0, row_idx), (5, row_idx), 'Helvetica-Bold')); page_style.append(('GRID', (0, row_idx), (5, row_idx), 0.5, COLOR_BORDER))
                    l_ptr += 1

                if row_idx <= right_occupied_until: pass
                elif r_ptr < len(right_queue):
                    item = right_queue[r_ptr]; itype = item['type']
                    if itype == 'title_section': row_content[7] = Paragraph(item['val'], section_title_style); page_style.append(('SPAN', (7, row_idx), (11, row_idx)))
                    elif itype == 'header_masuk':
                        cols = ['NO', 'TGL', 'SUMBER', 'JNS', 'LTR']
                        for c, txt in enumerate(cols): row_content[7+c] = Paragraph(txt, header_style)
                        page_style.append(('BACKGROUND', (7, row_idx), (11, row_idx), COLOR_HEADER_BLUE)); page_style.append(('GRID', (7, row_idx), (11, row_idx), 0.5, COLOR_BORDER))
                    elif itype == 'row_masuk': d = item['data']; row_content[7] = d[0]; row_content[8] = d[1]; row_content[9] = Paragraph(safe_text(d[2]), cell_style); row_content[10] = d[3]; row_content[11] = d[4]; page_style.append(('GRID', (7, row_idx), (11, row_idx), 0.5, COLOR_BORDER))
                    elif itype == 'total_masuk': row_content[7] = 'TOTAL MASUK'; row_content[11] = item['val']; page_style.append(('SPAN', (7, row_idx), (10, row_idx))); page_style.append(('BACKGROUND', (7, row_idx), (11, row_idx), COLOR_TOTAL_YELLOW)); page_style.append(('FONTNAME', (7, row_idx), (11, row_idx), 'Helvetica-Bold')); page_style.append(('GRID', (7, row_idx), (11, row_idx), 0.5, COLOR_BORDER))
                    elif itype == 'sub_rekap': style_to_use = header_black_style if item.get('txt_black') else header_style; row_content[7] = Paragraph(item['title'], style_to_use); bg = colors.HexColor(item['bg']); page_style.append(('SPAN', (7, row_idx), (11, row_idx))); page_style.append(('BACKGROUND', (7, row_idx), (11, row_idx), bg)); page_style.append(('GRID', (7, row_idx), (11, row_idx), 0.5, COLOR_BORDER))
                    elif itype == 'row_rekap': row_content[7] = Paragraph(safe_text(item['label'], 35), cell_style); row_content[11] = item['val']; page_style.append(('SPAN', (7, row_idx), (10, row_idx))); page_style.append(('GRID', (7, row_idx), (11, row_idx), 0.5, COLOR_BORDER))
                    elif itype == 'total_rekap': row_content[7] = 'TOTAL'; row_content[11] = item['val']; page_style.append(('SPAN', (7, row_idx), (10, row_idx))); page_style.append(('BACKGROUND', (7, row_idx), (11, row_idx), COLOR_TOTAL_YELLOW)); page_style.append(('FONTNAME', (7, row_idx), (11, row_idx), 'Helvetica-Bold')); page_style.append(('GRID', (7, row_idx), (11, row_idx), 0.5, COLOR_BORDER))
                    elif itype == 'header_stok': row_content[7] = Paragraph(item['label'], header_style); page_style.append(('SPAN', (7, row_idx), (11, row_idx))); page_style.append(('BACKGROUND', (7, row_idx), (11, row_idx), colors.HexColor("#70AD47"))); page_style.append(('GRID', (7, row_idx), (11, row_idx), 0.5, COLOR_BORDER))
                    elif itype == 'row_stok': row_content[7] = item['label']; row_content[11] = item['val']; page_style.append(('SPAN', (7, row_idx), (10, row_idx))); page_style.append(('GRID', (7, row_idx), (11, row_idx), 0.5, COLOR_BORDER))
                    elif itype == 'total_stok': row_content[7] = item['label']; row_content[11] = item['val']; page_style.append(('SPAN', (7, row_idx), (10, row_idx))); page_style.append(('BACKGROUND', (7, row_idx), (11, row_idx), colors.HexColor("#70AD47"))); page_style.append(('FONTNAME', (7, row_idx), (11, row_idx), 'Helvetica-Bold')); page_style.append(('TEXTCOLOR', (7, row_idx), (11, row_idx), colors.white)); page_style.append(('GRID', (7, row_idx), (11, row_idx), 0.5, COLOR_BORDER))
                    elif itype == 'chart':
                        span_needed = item['span']; rows_left = ROWS_PER_PAGE - row_idx
                        if rows_left < 5: pass
                        else:
                            real_span = min(span_needed, rows_left); img_height = real_span * 14
                            row_content[7] = RLImage(item['img'], width=200, height=img_height)
                            span_end_idx = row_idx + real_span - 1; page_style.append(('SPAN', (7, row_idx), (11, span_end_idx)))
                            right_occupied_until = span_end_idx; r_ptr += 1 
                    r_ptr += 1
                page_data.append(row_content); row_idx += 1
            if not page_data: break 
            col_widths = [20, 30, 80, 40, 30, 80,  20,  20, 30, 80, 40, 50]
            t = Table(page_data, colWidths=col_widths, rowHeights=[ROW_HEIGHT]*len(page_data)); t.setStyle(TableStyle(page_style)); elements.append(t)
            if l_ptr >= len(left_queue) and r_ptr >= len(right_queue): break
            elements.append(PageBreak())

    elements.append(PageBreak()); elements.append(Paragraph("LAPORAN BBM PERBULAN", title_style))
    m_data = []
    stok_run = hitung_stok_awal_periode(conn, lokasi_id, start_date_global)
    curr = start_date_global.replace(day=1); end_limit = end_date_global.replace(day=1)
    while curr <= end_limit:
        m = curr.month; y = curr.year
        q_in = f"SELECT SUM(jumlah_liter) FROM bbm_masuk WHERE lokasi_id={lokasi_id} AND MONTH(tanggal)={m} AND YEAR(tanggal)={y}"
        q_out = f"SELECT SUM(jumlah_liter) FROM bbm_keluar WHERE lokasi_id={lokasi_id} AND MONTH(tanggal)={m} AND YEAR(tanggal)={y}"
        cursor = conn.cursor()
        cursor.execute(q_in); res_in = cursor.fetchone(); mi = float(res_in[0]) if res_in and res_in[0] else 0.0
        cursor.execute(q_out); res_out = cursor.fetchone(); mo = float(res_out[0]) if res_out and res_out[0] else 0.0
        prev = stok_run; stok_run = prev + mi - mo
        m_data.append({'bln': curr.strftime("%B %Y"), 'awal': prev, 'masuk': mi, 'keluar': mo, 'sisa': stok_run, 'bulan_nama': get_bulan_indonesia(m)[:3]})
        curr = curr + relativedelta(months=1)
    
    df_m = pd.DataFrame(m_data)
    if not df_m.empty:
        img_m_buf = generate_monthly_chart(df_m)
        if img_m_buf: elements.append(RLImage(img_m_buf, width=480, height=220)); elements.append(Spacer(1, 15))
    
    d_m = [['BULAN', 'SISA BULAN LALU', 'MASUK', 'KELUAR', 'SISA']]
    for r in m_data: d_m.append([r['bln'], f"{r['awal']:,.0f}", f"{r['masuk']:,.0f}", f"{r['keluar']:,.0f}", f"{r['sisa']:,.0f}"])
    if m_data:
        t_masuk = sum(x['masuk'] for x in m_data); t_keluar = sum(x['keluar'] for x in m_data); akhir = m_data[-1]['sisa']
        d_m.append(['TOTAL', '', f"{t_masuk:,.0f}", f"{t_keluar:,.0f}", f"{akhir:,.0f}"])

    t_m = Table(d_m, colWidths=[100, 100, 100, 100, 100])
    rekap_style = [('GRID', (0,0), (-1,-1), 0.5, COLOR_BORDER), ('BACKGROUND', (0,0), (-1,0), COLOR_HEADER_BLUE), ('TEXTCOLOR', (0,0), (-1,0), colors.white), ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'), ('ALIGN', (0,0), (-1,0), 'CENTER'), ('ALIGN', (1,0), (-1,-1), 'RIGHT'), ('FONTSIZE', (0,0), (-1,-1), 9), ('LEFTPADDING', (0,0), (-1,-1), 6), ('RIGHTPADDING', (0,0), (-1,-1), 6)]
    for i in range(1, len(d_m)): bg = COLOR_ROW_EVEN if i % 2 == 0 else COLOR_ROW_ODD; rekap_style.append(('BACKGROUND', (0, i), (-1, i), bg))
    if m_data: rekap_style.append(('BACKGROUND', (0, -1), (-1, -1), COLOR_TOTAL_YELLOW)); rekap_style.append(('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'))
    t_m.setStyle(TableStyle(rekap_style)); elements.append(t_m)

    # --- ADD CHART TO PDF STANDARD (BELOW TABLE) ---
    df_keluar_all = pd.read_sql(f"SELECT * FROM bbm_keluar WHERE lokasi_id={lokasi_id} AND tanggal BETWEEN '{start_date_global}' AND '{end_date_global}'", conn)
    if not df_keluar_all.empty:
        if 'kategori' not in df_keluar_all.columns: df_keluar_all['kategori'] = df_keluar_all['nama_alat'].apply(cek_kategori)
        df_keluar_rpt = filter_non_consumption(df_keluar_all)
        df_alat_t, df_truck_t, _ = segregate_data(df_keluar_rpt, excluded_list)
        img_usage = generate_chart_for_report(df_alat_t, df_truck_t, width_inch=7, height_inch=3.5)
        if img_usage:
            elements.append(Spacer(1, 15))
            elements.append(RLImage(img_usage, width=480, height=240))

    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_docx_fixed(conn, lokasi_id, nama_lokasi, start_date_global, end_date_global, excluded_list):
    doc = Document(); 
    for s in doc.sections: s.left_margin=Cm(1); s.right_margin=Cm(1)
    date_ranges = split_date_range_by_month(start_date_global, end_date_global)

    for idx, (start_date, end_date) in enumerate(date_ranges):
        if idx > 0: doc.add_page_break()
        p = doc.add_paragraph(f"LAPORAN BBM: {nama_lokasi}"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold = True; p.runs[0].font.size = Pt(14)
        p2 = doc.add_paragraph(f"PERIODE {get_bulan_indonesia(start_date.month)} {start_date.year}"); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.runs[0].font.size = Pt(11)
        doc.add_paragraph()

        layout_table = doc.add_table(rows=1, cols=2); layout_table.autofit = False; layout_table.allow_autofit = False
        layout_table.columns[0].width = Cm(10); layout_table.columns[1].width = Cm(9)
        cell_left = layout_table.cell(0, 0); cell_right = layout_table.cell(0, 1)

        stok_awal = hitung_stok_awal_periode(conn, lokasi_id, start_date)
        df_keluar = pd.read_sql(f"SELECT * FROM bbm_keluar WHERE lokasi_id={lokasi_id} AND tanggal BETWEEN '{start_date}' AND '{end_date}' ORDER BY tanggal", conn)
        if not df_keluar.empty and 'kategori' not in df_keluar.columns: df_keluar['kategori'] = df_keluar['nama_alat'].apply(cek_kategori)
        df_keluar_rpt = filter_non_consumption(df_keluar)
        
        cell_left.add_paragraph("PENGGUNAAN BBM (KELUAR)", style='Heading 3')
        tbl_k = cell_left.add_table(rows=1, cols=4); tbl_k.style = 'Table Grid'
        h_k = tbl_k.rows[0].cells; h_k[0].text="TGL"; h_k[1].text="ALAT"; h_k[2].text="UNIT"; h_k[3].text="LTR"
        if not df_keluar_rpt.empty:
            last_date = None; is_grey = False
            for i, r in df_keluar_rpt.iterrows():
                curr_date = r['tanggal']; is_grey = not is_grey if last_date is not None and curr_date != last_date else is_grey; last_date = curr_date
                row = tbl_k.add_row().cells
                if is_grey:
                    for c in row: set_cell_bg(c, "F2F2F2")
                row[0].text = r['tanggal'].strftime('%d/%m'); row[1].text = r['nama_alat']; row[2].text = r['no_unit']; row[3].text = f"{r['jumlah_liter']:.0f}"
                for c in row: c.paragraphs[0].runs[0].font.size = Pt(8)
        
        df_masuk = pd.read_sql(f"SELECT * FROM bbm_masuk WHERE lokasi_id={lokasi_id} AND tanggal BETWEEN '{start_date}' AND '{end_date}' ORDER BY tanggal", conn)
        cell_right.add_paragraph("BBM MASUK", style='Heading 3')
        tbl_m = cell_right.add_table(rows=1, cols=3); tbl_m.style='Table Grid'
        h_m = tbl_m.rows[0].cells; h_m[0].text="TGL"; h_m[1].text="SUMBER"; h_m[2].text="LTR"
        if not df_masuk.empty:
            for i, r in df_masuk.iterrows():
                row = tbl_m.add_row().cells; row[0].text = r['tanggal'].strftime('%d/%m'); row[1].text = r['sumber']; row[2].text = f"{r['jumlah_liter']:.0f}"
                for c in row: c.paragraphs[0].runs[0].font.size = Pt(8)
        cell_right.add_paragraph("")
        cell_right.add_paragraph("RINCIAN PENGGUNAAN BBM", style='Heading 4')
        df_alat_g, df_truck_g, df_lain_g = segregate_data(df_keluar_rpt, excluded_list)
        
        def add_detailed_docx(container, title, df_subset, color_hex):
            p = container.add_paragraph(title); p.runs[0].font.bold=True; p.runs[0].font.size=Pt(9)
            t = container.add_table(rows=1, cols=2); t.style='Table Grid'; total_liter = 0
            if not df_subset.empty and 'jumlah_liter' in df_subset.columns:
                total_liter = df_subset['jumlah_liter'].sum(); grp = df_subset.groupby(['nama_alat', 'no_unit'])['jumlah_liter'].sum().reset_index()
                for _, r in grp.iterrows(): row = t.add_row().cells; row[0].text = f"{r['nama_alat']} {r['no_unit']}"; row[1].text = f"{r['jumlah_liter']:.0f}"; 
                for c in row: c.paragraphs[0].runs[0].font.size = Pt(8)
            else: t.add_row().cells[0].text = "KOSONG"
            rt = t.add_row().cells; rt[0].text="TOTAL"; rt[1].text=f"{total_liter:.0f}"
            for c in rt: 
                set_cell_bg(c, "FFFF00")
                if len(c.paragraphs)>0 and len(c.paragraphs[0].runs)>0: c.paragraphs[0].runs[0].bold=True
                elif len(c.paragraphs)>0: c.paragraphs[0].add_run(c.text).font.bold = True
        add_detailed_docx(cell_right, "TOTAL PENGGUNAAN BBM ALAT BERAT", df_alat_g, "F4B084"); cell_right.add_paragraph(""); add_detailed_docx(cell_right, "TOTAL PENGGUNAAN BBM MOBIL & TRUCK", df_truck_g, "9BC2E6"); cell_right.add_paragraph("")
        if not df_lain_g.empty: add_detailed_docx(cell_right, "TOTAL PENGGUNAAN BBM LAINNYA", df_lain_g, "FFB6C1"); cell_right.add_paragraph("")

        tm = float(df_masuk['jumlah_liter'].sum()) if not df_masuk.empty else 0.0
        tk_real = float(df_keluar['jumlah_liter'].sum()) if not df_keluar.empty else 0.0
        sisa = stok_awal + tm - tk_real
        cell_right.add_paragraph("RINCIAN SISA STOK BBM", style='Heading 4')
        tbl_s = cell_right.add_table(rows=4, cols=2); tbl_s.style='Table Grid'
        tbl_s.cell(0,0).text="SISA BULAN LALU"; tbl_s.cell(0,1).text=f"{stok_awal:.0f}"
        tbl_s.cell(1,0).text="TOTAL MASUK"; tbl_s.cell(1,1).text=f"{tm:.0f}"
        tbl_s.cell(2,0).text="TOTAL KELUAR (REAL)"; tbl_s.cell(2,1).text=f"{tk_real:.0f}"
        tbl_s.cell(3,0).text="SISA AKHIR"; tbl_s.cell(3,1).text=f"{sisa:.0f}"
        
        img_buf = generate_chart_for_report(df_alat_g, df_truck_g, width_inch=3.5, height_inch=2.5)
        if img_buf: 
            cell_right.add_paragraph("")
            cell_right.add_paragraph().add_run().add_picture(img_buf, width=Cm(8))

    doc.add_page_break(); p_title = doc.add_paragraph("LAPORAN BBM PERBULAN"); p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_title.runs[0].bold=True; p_title.runs[0].font.size=Pt(14)
    m_data = []; stok_run = hitung_stok_awal_periode(conn, lokasi_id, start_date_global)
    curr = start_date_global.replace(day=1); end_limit = end_date_global.replace(day=1)
    while curr <= end_limit:
        m = curr.month; y = curr.year
        q_in = f"SELECT SUM(jumlah_liter) FROM bbm_masuk WHERE lokasi_id={lokasi_id} AND MONTH(tanggal)={m} AND YEAR(tanggal)={y}"
        q_out = f"SELECT SUM(jumlah_liter) FROM bbm_keluar WHERE lokasi_id={lokasi_id} AND MONTH(tanggal)={m} AND YEAR(tanggal)={y}"
        cursor = conn.cursor()
        cursor.execute(q_in); res_in = cursor.fetchone(); mi = float(res_in[0]) if res_in and res_in[0] else 0.0
        cursor.execute(q_out); res_out = cursor.fetchone(); mo = float(res_out[0]) if res_out and res_out[0] else 0.0
        prev = stok_run; stok_run = prev + mi - mo
        m_data.append({'bln': curr.strftime("%B %Y"), 'awal': prev, 'masuk': mi, 'keluar': mo, 'sisa': stok_run, 'bulan_nama': get_bulan_indonesia(m)[:3]})
        curr = curr + relativedelta(months=1)

    df_m = pd.DataFrame(m_data)
    if not df_m.empty:
        img_m_buf = generate_monthly_chart(df_m)
        if img_m_buf: doc.add_paragraph().add_run().add_picture(img_m_buf, width=Cm(16))
    doc.add_paragraph("RINCIAN MASUK DAN PENGGUNAAN SOLAR PERBULANNYA", style='Heading 4')
    tbl_month = doc.add_table(rows=1, cols=5); tbl_month.style='Table Grid'
    h_month = tbl_month.rows[0].cells
    for i, t in enumerate(['BULAN', 'SISA BULAN LALU', 'MASUK', 'KELUAR', 'SISA']): h_month[i].text=t
    for r in m_data:
        row = tbl_month.add_row().cells; row[0].text=r['bln']; row[1].text=f"{r['awal']:.0f}"; row[2].text=f"{r['masuk']:.0f}"; row[3].text=f"{r['keluar']:.0f}"; row[4].text=f"{r['sisa']:.0f}"
    
    if m_data:
        t_masuk = sum(x['masuk'] for x in m_data); t_keluar = sum(x['keluar'] for x in m_data); akhir = m_data[-1]['sisa']
        row = tbl_month.add_row().cells
        row[0].text = "TOTAL"; row[2].text = f"{t_masuk:,.0f}"; row[3].text = f"{t_keluar:,.0f}"; row[4].text = f"{akhir:,.0f}"
        for c in row:
            set_cell_bg(c, "FFD966")
            if len(c.paragraphs) > 0 and len(c.paragraphs[0].runs) > 0: c.paragraphs[0].runs[0].font.bold = True
            elif len(c.paragraphs) > 0: c.paragraphs[0].add_run(c.text).font.bold = True

    # --- ADD CHART TO DOCX STANDARD (BELOW TABLE) ---
    df_keluar_all = pd.read_sql(f"SELECT * FROM bbm_keluar WHERE lokasi_id={lokasi_id} AND tanggal BETWEEN '{start_date_global}' AND '{end_date_global}'", conn)
    if not df_keluar_all.empty:
        if 'kategori' not in df_keluar_all.columns: df_keluar_all['kategori'] = df_keluar_all['nama_alat'].apply(cek_kategori)
        df_keluar_rpt = filter_non_consumption(df_keluar_all)
        df_alat_t, df_truck_t, _ = segregate_data(df_keluar_rpt, excluded_list)
        img_usage = generate_chart_for_report(df_alat_t, df_truck_t, width_inch=7, height_inch=3.5)
        if img_usage:
            doc.add_paragraph().add_run().add_picture(img_usage, width=Cm(16))

    buffer = io.BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer

def generate_excel_styled(conn, lokasi_id, nama_lokasi, start_date_global, end_date_global, excluded_list):
    output = io.BytesIO(); wb = Workbook(); wb.remove(wb.active)
    thin = Border(left=Side('thin'), right=Side('thin'), top=Side('thin'), bottom=Side('thin'))
    date_ranges = split_date_range_by_month(start_date_global, end_date_global)
    for idx, (start_date, end_date) in enumerate(date_ranges):
        sheet_name = get_bulan_indonesia(start_date.month)[:3] + f" {start_date.year}"
        ws = wb.create_sheet(sheet_name)
        ws.column_dimensions['A'].width = 5; ws.column_dimensions['B'].width = 15; ws.column_dimensions['C'].width = 30; ws.column_dimensions['D'].width = 15
        ws.column_dimensions['E'].width = 15; ws.column_dimensions['F'].width = 35
        ws.column_dimensions['I'].width = 5; ws.column_dimensions['J'].width = 15; ws.column_dimensions['K'].width = 30; ws.column_dimensions['L'].width = 15
        
        stok_awal = hitung_stok_awal_periode(conn, lokasi_id, start_date)
        df_masuk = pd.read_sql(f"SELECT * FROM bbm_masuk WHERE lokasi_id={lokasi_id} AND tanggal BETWEEN '{start_date}' AND '{end_date}' ORDER BY tanggal", conn)
        df_keluar = pd.read_sql(f"SELECT * FROM bbm_keluar WHERE lokasi_id={lokasi_id} AND tanggal BETWEEN '{start_date}' AND '{end_date}' ORDER BY tanggal", conn)
        
        if not df_keluar.empty and 'kategori' not in df_keluar.columns: df_keluar['kategori'] = df_keluar['nama_alat'].apply(cek_kategori)
        df_keluar_rpt = filter_non_consumption(df_keluar)
        tm = float(df_masuk['jumlah_liter'].sum()) if not df_masuk.empty else 0.0
        tk_real = float(df_keluar['jumlah_liter'].sum()) if not df_keluar.empty else 0.0
        tk_rpt = float(df_keluar_rpt['jumlah_liter'].sum()) if not df_keluar_rpt.empty else 0.0
        sisa_akhir = stok_awal + tm - tk_real
        
        ws.merge_cells('A1:L1'); ws['A1'] = f"LAPORAN BBM: {nama_lokasi}"; ws['A1'].font = Font(bold=True, size=14); ws['A1'].alignment = Alignment(horizontal='center')
        ws.merge_cells('A2:L2'); ws['A2'] = f"PERIODE {get_bulan_indonesia(start_date.month)} {start_date.year}"; ws['A2'].font = Font(size=12); ws['A2'].alignment = Alignment(horizontal='center')
        
        r = 4; ws.cell(r, 1, "PENGGUNAAN BBM (KELUAR)").font = Font(bold=True)
        headers = ['NO', 'TGL', 'ALAT', 'UNIT', 'LTR', 'KET']
        for i, h in enumerate(headers): c=ws.cell(r+1, i+1, h); c.border=thin; c.fill=PatternFill("solid", fgColor="D3D3D3"); c.alignment=Alignment(horizontal='center')
        r += 2
        if not df_keluar_rpt.empty:
            last_date = None; is_grey = False
            for i, row in df_keluar_rpt.iterrows():
                curr_date = row['tanggal']
                if last_date is not None and curr_date != last_date: is_grey = not is_grey
                last_date = curr_date
                fill = PatternFill("solid", fgColor="F2F2F2") if is_grey else None
                vals = [i+1, row['tanggal'].strftime('%d/%m/%Y'), row['nama_alat'], row['no_unit'], float(row['jumlah_liter']), row['keterangan']]
                for j, v in enumerate(vals): c=ws.cell(r, j+1, v); c.border=thin; c.alignment=Alignment(wrap_text=True, vertical='center'); 
                if fill: 
                    for j in range(6): ws.cell(r, j+1).fill = fill
                r += 1
        ws.cell(r, 3, "TOTAL").font=Font(bold=True); c=ws.cell(r, 5, tk_rpt); c.font=Font(bold=True); c.fill=PatternFill("solid", fgColor="FFFF00"); c.border=thin
        
        r_r = 4; ws.cell(r_r, 9, "BBM MASUK").font = Font(bold=True)
        headers_m = ['NO', 'TGL', 'SUMBER', 'JNS', 'LTR']
        for i, h in enumerate(headers_m): c=ws.cell(r_r+1, i+9, h); c.border=thin; c.fill=PatternFill("solid", fgColor="D3D3D3"); c.alignment=Alignment(horizontal='center')
        r_r += 2
        if not df_masuk.empty:
            for i, row in df_masuk.iterrows():
                vals = [i+1, row['tanggal'].strftime('%d/%m/%Y'), row['sumber'], row['jenis_bbm'], float(row['jumlah_liter'])]
                for j, v in enumerate(vals): c=ws.cell(r_r, j+9, v); c.border=thin; c.alignment=Alignment(wrap_text=True)
                r_r += 1
        ws.cell(r_r, 11, "TOTAL").font=Font(bold=True); c=ws.cell(r_r, 13, tm); c.font=Font(bold=True); c.fill=PatternFill("solid", fgColor="FFFF00"); c.border=thin
        r_r += 2
        
        ws.cell(r_r, 9, "RINCIAN PENGGUNAAN BBM").font=Font(bold=True); r_r+=1
        df_alat_g, df_truck_g, df_lain_g = segregate_data(df_keluar_rpt, excluded_list)
        
        def write_detail(ws, row, col, title, df, color):
            ws.merge_cells(start_row=row, start_column=col, end_row=row, end_column=col+3)
            c=ws.cell(row, col, title); c.fill=color; c.font=Font(bold=True); c.alignment=Alignment(horizontal='center'); c.border=thin
            row+=1
            if not df.empty and 'jumlah_liter' in df.columns:
                grp = df.groupby(['nama_alat', 'no_unit'])['jumlah_liter'].sum().reset_index()
                for _, x in grp.iterrows():
                    ws.merge_cells(start_row=row, start_column=col, end_row=row, end_column=col+2)
                    c1=ws.cell(row, col, f"{x['nama_alat']} {x['no_unit']}"); c1.border=thin
                    c2=ws.cell(row, col+3, float(x['jumlah_liter'])); c2.border=thin
                    row+=1
            total_val = float(df['jumlah_liter'].sum()) if not df.empty and 'jumlah_liter' in df.columns else 0
            ws.merge_cells(start_row=row, start_column=col, end_row=row, end_column=col+2)
            c_tot=ws.cell(row, col, "TOTAL"); c_tot.fill=PatternFill("solid", fgColor="FFFF00"); c_tot.border=thin; c_tot.font=Font(bold=True)
            c_val=ws.cell(row, col+3, total_val); c_val.fill=PatternFill("solid", fgColor="FFFF00"); c_val.border=thin; c_val.font=Font(bold=True)
            return row+2
        
        r_r = write_detail(ws, r_r, 9, "TOTAL PENGGUNAAN ALAT BERAT", df_alat_g, PatternFill("solid", fgColor="F4B084"))
        r_r = write_detail(ws, r_r, 9, "TOTAL PENGGUNAAN MOBIL & TRUCK", df_truck_g, PatternFill("solid", fgColor="9BC2E6"))
        if not df_lain_g.empty: r_r = write_detail(ws, r_r, 9, "TOTAL PENGGUNAAN BBM LAINNYA", df_lain_g, PatternFill("solid", fgColor="FFB6C1"))
            
        ws.cell(r_r, 9, "RINCIAN SISA STOK BBM").font=Font(bold=True); r_r+=1
        data_s = [('SISA BULAN LALU', stok_awal), ('MASUK', tm), ('KELUAR (REAL)', tk_real), ('SISA AKHIR', sisa_akhir)]
        for k, v in data_s:
            ws.merge_cells(start_row=r_r, start_column=9, end_row=r_r, end_column=11)
            c1=ws.cell(r_r, 9, k); c1.border=thin
            c2=ws.cell(r_r, 12, v); c2.border=thin
            if k == 'SISA AKHIR': c1.fill=PatternFill("solid", fgColor="00FF00"); c2.fill=PatternFill("solid", fgColor="00FF00")
            r_r+=1
        r_r+=1
        
        img_buf = generate_chart_for_report(df_alat_g, df_truck_g, width_inch=4.5, height_inch=3.0)
        if img_buf: 
            img = XLImage(img_buf); img.width = 450; img.height = 450
            ws.add_image(img, f'I{r_r}')

    ws2 = wb.create_sheet("Rekap Tahunan"); ws2['A1'] = "LAPORAN BBM PERBULAN"; ws2['A1'].font = Font(bold=True, size=14)
    ws2.column_dimensions['A'].width = 25; ws2.column_dimensions['B'].width = 20; ws2.column_dimensions['C'].width = 20; ws2.column_dimensions['D'].width = 20; ws2.column_dimensions['E'].width = 20
    m_data = []
    stok_run = hitung_stok_awal_periode(conn, lokasi_id, start_date_global)
    curr = start_date_global.replace(day=1); end_limit = end_date_global.replace(day=1)
    while curr <= end_limit:
        m = curr.month; y = curr.year
        q_in = f"SELECT SUM(jumlah_liter) FROM bbm_masuk WHERE lokasi_id={lokasi_id} AND MONTH(tanggal)={m} AND YEAR(tanggal)={y}"
        q_out = f"SELECT SUM(jumlah_liter) FROM bbm_keluar WHERE lokasi_id={lokasi_id} AND MONTH(tanggal)={m} AND YEAR(tanggal)={y}"
        cursor = conn.cursor()
        cursor.execute(q_in); res_in = cursor.fetchone(); mi = float(res_in[0]) if res_in and res_in[0] else 0.0
        cursor.execute(q_out); res_out = cursor.fetchone(); mo = float(res_out[0]) if res_out and res_out[0] else 0.0
        prev = stok_run; stok_run = prev + mi - mo
        m_data.append({'bln': curr.strftime("%B %Y"), 'awal': prev, 'masuk': mi, 'keluar': mo, 'sisa': stok_run, 'bulan_nama': get_bulan_indonesia(m)[:3]})
        curr = curr + relativedelta(months=1)
    df_m = pd.DataFrame(m_data)
    img_m_buf = generate_monthly_chart(df_m)
    if img_m_buf: img2 = XLImage(img_m_buf); img2.width=500; img2.height=250; ws2.add_image(img2, 'A3')
    r2 = 18; headers = ['BULAN', 'SISA BULAN LALU', 'MASUK', 'KELUAR', 'SISA']
    for i, h in enumerate(headers): c=ws2.cell(r2, i+1, h); c.border=thin; c.fill=PatternFill("solid", fgColor="D3D3D3")
    r2+=1
    for r in m_data:
        vals = [r['bln'], r['awal'], r['masuk'], r['keluar'], r['sisa']]
        for i, v in enumerate(vals): c=ws2.cell(r2, i+1, v); c.border=thin
        r2+=1
    if m_data:
        t_masuk = sum(x['masuk'] for x in m_data); t_keluar = sum(x['keluar'] for x in m_data); akhir = m_data[-1]['sisa']
        ws2.cell(r2, 1, "TOTAL").font = Font(bold=True); ws2.cell(r2, 3, t_masuk).font = Font(bold=True); ws2.cell(r2, 4, t_keluar).font = Font(bold=True); ws2.cell(r2, 5, akhir).font = Font(bold=True)
        for i in range(1, 6): c = ws2.cell(r2, i); c.fill = PatternFill("solid", fgColor="FFD966"); c.border = thin
    
    # --- ADD CHART TO EXCEL STANDARD (RIGHT OF MONTHLY CHART) ---
    df_keluar_all = pd.read_sql(f"SELECT * FROM bbm_keluar WHERE lokasi_id={lokasi_id} AND tanggal BETWEEN '{start_date_global}' AND '{end_date_global}'", conn)
    if not df_keluar_all.empty:
        if 'kategori' not in df_keluar_all.columns: df_keluar_all['kategori'] = df_keluar_all['nama_alat'].apply(cek_kategori)
        df_keluar_rpt = filter_non_consumption(df_keluar_all)
        df_alat_t, df_truck_t, _ = segregate_data(df_keluar_rpt, excluded_list)
        img_usage = generate_chart_for_report(df_alat_t, df_truck_t, width_inch=7, height_inch=3.5)
        if img_usage:
            img3 = XLImage(img_usage); img3.width=500; img3.height=250
            ws2.add_image(img3, 'H3') # Positioned at H3, roughly to the right of A3

    wb.save(output); output.seek(0)
    return output

def generate_pdf_one_sheet(conn, lokasi_id, nama_lokasi, start_date_global, end_date_global, excluded_list):
    buffer = io.BytesIO()
    
    date_ranges = split_date_range_by_month(start_date_global, end_date_global)
    
    SPLIT_IDX = 115
    # UPDATE FIX: Calculate page height dynamically based on content volume to prevent LayoutError
    # Base height + (rows * height) + buffer for charts
    page_height = max(2000, (SPLIT_IDX * 20) + 1200) 
    page_width = 35 * cm 
    
    doc = SimpleDocTemplate(buffer, pagesize=(page_width, page_height), rightMargin=20, leftMargin=20, topMargin=20, bottomMargin=20)
    elements = []
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(name='Title', parent=styles['Heading1'], alignment=TA_CENTER, fontSize=16, fontName='Helvetica-Bold', spaceAfter=2)
    period_style = ParagraphStyle(name='Period', parent=styles['Normal'], alignment=TA_CENTER, fontSize=12, fontName='Helvetica-Bold', spaceAfter=15)
    header_style = ParagraphStyle(name='Header', parent=styles['Normal'], fontSize=8, fontName='Helvetica-Bold', textColor=colors.white, alignment=TA_CENTER)
    cell_style = ParagraphStyle(name='Cell', parent=styles['Normal'], fontSize=8, fontName='Helvetica')
    h3_style = ParagraphStyle(name='H3', parent=styles['Heading3'], fontSize=10, fontName='Helvetica-Bold', spaceAfter=4)

    for idx, (start_date, end_date) in enumerate(date_ranges):
        if idx > 0: elements.append(PageBreak())
        
        stok_awal = hitung_stok_awal_periode(conn, lokasi_id, start_date)
        df_masuk = pd.read_sql(f"SELECT * FROM bbm_masuk WHERE lokasi_id={lokasi_id} AND tanggal BETWEEN '{start_date}' AND '{end_date}' ORDER BY tanggal", conn)
        df_keluar = pd.read_sql(f"SELECT * FROM bbm_keluar WHERE lokasi_id={lokasi_id} AND tanggal BETWEEN '{start_date}' AND '{end_date}' ORDER BY tanggal", conn)
        
        if not df_keluar.empty and 'kategori' not in df_keluar.columns: df_keluar['kategori'] = df_keluar['nama_alat'].apply(cek_kategori)
        df_keluar_rpt = filter_non_consumption(df_keluar)
        df_alat_g, df_truck_g, df_lain_g = segregate_data(df_keluar_rpt, excluded_list)
        
        tm = float(df_masuk['jumlah_liter'].sum()) if not df_masuk.empty else 0.0
        tk_real = float(df_keluar['jumlah_liter'].sum()) if not df_keluar.empty else 0.0
        tk_rpt = float(df_keluar_rpt['jumlah_liter'].sum()) if not df_keluar_rpt.empty else 0.0
        sisa_akhir = stok_awal + tm - tk_real
        
        df_left = df_keluar_rpt.iloc[:SPLIT_IDX]
        df_right_top = df_keluar_rpt.iloc[SPLIT_IDX:]
        
        elements.append(Paragraph(f"LAPORAN BBM: {nama_lokasi}", title_style))
        elements.append(Paragraph(f"PERIODE {get_bulan_indonesia(start_date.month)} {start_date.year}", period_style))

        # --- LEFT COLUMN TABLE ---
        left_stack = []
        left_stack.append(Paragraph("PENGGUNAAN BBM (KELUAR)", h3_style))
        table_left_data = [['NO', 'TGL', 'ALAT', 'UNIT', 'LTR', 'KET']]
        
        left_row_bg = []
        last_date = None; is_grey = False
        row_count = 0
        for i, r in df_left.iterrows():
            curr_date = r['tanggal']
            if last_date is not None and curr_date != last_date: is_grey = not is_grey
            last_date = curr_date
            if is_grey: left_row_bg.append(row_count + 1)
            table_left_data.append([i+1, r['tanggal'].strftime('%d/%m'), Paragraph(safe_text(r['nama_alat']), cell_style), Paragraph(safe_text(r['no_unit']), cell_style), f"{r['jumlah_liter']:.0f}", Paragraph(safe_text(r['keterangan']), cell_style)])
            row_count += 1
        
        if df_right_top.empty:
            table_left_data.append(['', 'TOTAL', '', '', f"{tk_rpt:.0f}", ''])

        t_left = Table(table_left_data, colWidths=[25, 40, 100, 50, 40, 100])
        style_left = [
            ('GRID', (0,0), (-1,-1), 0.5, COLOR_BORDER),
            ('BACKGROUND', (0,0), (-1,0), COLOR_HEADER_BLUE),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('ALIGN', (0,0), (-1,0), 'CENTER'),
            ('ALIGN', (4,1), (4,-1), 'RIGHT'),
            ('FONTSIZE', (0,0), (-1,-1), 8),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]
        for rid in left_row_bg: style_left.append(('BACKGROUND', (0, rid), (-1, rid), COLOR_ROW_EVEN))
        
        if df_right_top.empty:
             style_left.append(('BACKGROUND', (0,-1), (-1,-1), COLOR_TOTAL_YELLOW))
             style_left.append(('FONTNAME', (0,-1), (-1,-1), 'Helvetica-Bold'))
             style_left.append(('SPAN', (1,-1), (2,-1)))

        t_left.setStyle(TableStyle(style_left))
        left_stack.append(t_left)
        
        # --- RIGHT COLUMN STACK ---
        right_stack = []
        
        if not df_right_top.empty:
            right_stack.append(Paragraph("PENGGUNAAN BBM (LANJUTAN)", h3_style))
            table_right_top_data = [['NO', 'TGL', 'ALAT', 'UNIT', 'LTR', 'KET']]
            rt_row_bg = []
            row_count_rt = 0
            last_date_rt = None; is_grey_rt = False
            for i, r in df_right_top.iterrows():
                curr_date = r['tanggal']
                if last_date_rt is not None and curr_date != last_date_rt: is_grey_rt = not is_grey_rt
                last_date_rt = curr_date
                if is_grey_rt: rt_row_bg.append(row_count_rt + 1)
                table_right_top_data.append([SPLIT_IDX + i + 1, r['tanggal'].strftime('%d/%m'), Paragraph(safe_text(r['nama_alat']), cell_style), Paragraph(safe_text(r['no_unit']), cell_style), f"{r['jumlah_liter']:.0f}", Paragraph(safe_text(r['keterangan']), cell_style)])
                row_count_rt += 1
            table_right_top_data.append(['', 'TOTAL', '', '', f"{tk_rpt:.0f}", ''])

            t_rt = Table(table_right_top_data, colWidths=[25, 40, 100, 50, 40, 100])
            style_rt = [
                ('GRID', (0,0), (-1,-1), 0.5, COLOR_BORDER),
                ('BACKGROUND', (0,0), (-1,0), COLOR_HEADER_BLUE),
                ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('ALIGN', (0,0), (-1,0), 'CENTER'),
                ('ALIGN', (4,1), (4,-1), 'RIGHT'),
                ('FONTSIZE', (0,0), (-1,-1), 8),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('BACKGROUND', (0,-1), (-1,-1), COLOR_TOTAL_YELLOW), # Total row
                ('FONTNAME', (0,-1), (-1,-1), 'Helvetica-Bold'),
                ('SPAN', (1,-1), (2,-1))
            ]
            for rid in rt_row_bg: style_rt.append(('BACKGROUND', (0, rid), (-1, rid), COLOR_ROW_EVEN))
            t_rt.setStyle(TableStyle(style_rt))
            right_stack.append(t_rt)
            right_stack.append(Spacer(1, 5))

        right_stack.append(Paragraph("BBM MASUK", h3_style))
        table_masuk_data = [['NO', 'TGL', 'SUMBER', 'JNS', 'LTR']]
        if not df_masuk.empty:
            for i, r in df_masuk.iterrows(): table_masuk_data.append([i+1, r['tanggal'].strftime('%d/%m'), r['sumber'], r['jenis_bbm'], f"{r['jumlah_liter']:.0f}"])
        else: table_masuk_data.append(['-', '-', 'TIDAK ADA', '-', '0'])
        table_masuk_data.append(['', 'TOTAL', '', '', f"{tm:.0f}"])
        
        t_masuk = Table(table_masuk_data, colWidths=[25, 40, 120, 50, 60])
        t_masuk.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, COLOR_BORDER),
            ('BACKGROUND', (0,0), (-1,0), COLOR_HEADER_BLUE),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('ALIGN', (0,0), (-1,0), 'CENTER'),
            ('BACKGROUND', (0,-1), (-1,-1), COLOR_TOTAL_YELLOW), # Total row
            ('FONTNAME', (0,-1), (-1,-1), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,-1), 8),
            ('SPAN', (1,-1), (2,-1))
        ]))
        right_stack.append(t_masuk)
        right_stack.append(Spacer(1, 5))
        
        right_stack.append(Paragraph("RINCIAN PENGGUNAAN BBM", h3_style))
        def create_rekap_table(title, df_subset, color_hex):
            if df_subset.empty: return None
            data = [[title, '', '']]
            grp = df_subset.groupby(['nama_alat', 'no_unit'])['jumlah_liter'].sum().reset_index().sort_values('jumlah_liter', ascending=False)
            for _, r in grp.iterrows(): data.append([r['nama_alat'], r['no_unit'], f"{r['jumlah_liter']:.0f}"])
            data.append(['TOTAL', '', f"{df_subset['jumlah_liter'].sum():.0f}"])
            t = Table(data, colWidths=[110, 80, 60])
            s = [('GRID', (0,0), (-1,-1), 0.5, COLOR_BORDER),('BACKGROUND', (0,0), (-1,0), colors.HexColor(color_hex)),('SPAN', (0,0), (-1,0)),('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),('BACKGROUND', (0,-1), (-1,-1), COLOR_TOTAL_YELLOW),('FONTNAME', (0,-1), (-1,-1), 'Helvetica-Bold'),('FONTSIZE', (0,0), (-1,-1), 8),]
            t.setStyle(TableStyle(s)); return t

        t_ab = create_rekap_table("TOTAL ALAT BERAT", df_alat_g, "#F4B084")
        if t_ab: right_stack.append(t_ab); right_stack.append(Spacer(1, 2))
        
        t_mt = create_rekap_table("TOTAL MOBIL & TRUCK", df_truck_g, "#9BC2E6")
        if t_mt: right_stack.append(t_mt); right_stack.append(Spacer(1, 2))

        t_ot = create_rekap_table("LAINNYA", df_lain_g, "#ED77C4")
        if t_ot: right_stack.append(t_ot); right_stack.append(Spacer(1, 2))

        right_stack.append(Spacer(1, 5))
        right_stack.append(Paragraph("RINCIAN SISA STOK BBM", h3_style))
        stok_data = [['RINGKASAN STOK', ''], ['SISA BULAN LALU', f"{stok_awal:.0f}"], ['TOTAL MASUK', f"{tm:.0f}"], ['TOTAL KELUAR', f"{tk_real:.0f}"], ['SISA AKHIR', f"{sisa_akhir:.0f}"]]
        t_stok = Table(stok_data, colWidths=[150, 80])
        t_stok.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 0.5, COLOR_BORDER),('BACKGROUND', (0,0), (-1,0), colors.HexColor("#70AD47")),('BACKGROUND', (0,4), (-1,4), colors.HexColor("#00FF00")),('FONTSIZE', (0,0), (-1,-1), 8)]))
        right_stack.append(t_stok)
        
        img_buf = generate_chart_for_report(df_alat_g, df_truck_g, width_inch=3.5, height_inch=2.5)
        if img_buf: 
            right_stack.append(Spacer(1, 5))
            right_stack.append(RLImage(img_buf, width=200, height=200))
        
        main_table_data = [[left_stack, right_stack]]
        t_main = Table(main_table_data, colWidths=[380, 400], vAlign='TOP')
        t_main.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'TOP'), ('LEFTPADDING', (0,0), (-1,-1), 0), ('RIGHTPADDING', (0,0), (-1,-1), 0)]))
        elements.append(t_main)
    
    elements.append(PageBreak())
    elements.append(Paragraph("LAPORAN BBM PERBULAN", title_style))
    m_data = []; stok_run = hitung_stok_awal_periode(conn, lokasi_id, start_date_global)
    curr = start_date_global.replace(day=1); end_limit = end_date_global.replace(day=1)
    while curr <= end_limit:
        m = curr.month; y = curr.year
        q_in = f"SELECT SUM(jumlah_liter) FROM bbm_masuk WHERE lokasi_id={lokasi_id} AND MONTH(tanggal)={m} AND YEAR(tanggal)={y}"
        q_out = f"SELECT SUM(jumlah_liter) FROM bbm_keluar WHERE lokasi_id={lokasi_id} AND MONTH(tanggal)={m} AND YEAR(tanggal)={y}"
        cursor = conn.cursor()
        cursor.execute(q_in); res_in = cursor.fetchone(); mi = float(res_in[0]) if res_in and res_in[0] else 0.0
        cursor.execute(q_out); res_out = cursor.fetchone(); mo = float(res_out[0]) if res_out and res_out[0] else 0.0
        prev = stok_run; stok_run = prev + mi - mo
        m_data.append({'bln': curr.strftime("%B %Y"), 'awal': prev, 'masuk': mi, 'keluar': mo, 'sisa': stok_run, 'bulan_nama': get_bulan_indonesia(m)[:3]})
        curr = curr + relativedelta(months=1)

    df_m = pd.DataFrame(m_data)
    
    # --- CHART LOGIC FOR PDF ONE SHEET (RIGHT SIDE) ---
    img_m_buf = None
    if not df_m.empty:
        img_m_buf = generate_monthly_chart(df_m)
    
    # Generate Usage Chart for the whole period
    df_keluar_all = pd.read_sql(f"SELECT * FROM bbm_keluar WHERE lokasi_id={lokasi_id} AND tanggal BETWEEN '{start_date_global}' AND '{end_date_global}'", conn)
    img_usage_buf = None
    if not df_keluar_all.empty:
        if 'kategori' not in df_keluar_all.columns: df_keluar_all['kategori'] = df_keluar_all['nama_alat'].apply(cek_kategori)
        df_keluar_rpt = filter_non_consumption(df_keluar_all)
        df_alat_t, df_truck_t, _ = segregate_data(df_keluar_rpt, excluded_list)
        img_usage_buf = generate_chart_for_report(df_alat_t, df_truck_t, width_inch=7, height_inch=3.5)

    # Layout for charts: Side by Side if both exist
    chart_row = []
    if img_m_buf: chart_row.append(RLImage(img_m_buf, width=400, height=200))
    else: chart_row.append("")
    if img_usage_buf: chart_row.append(RLImage(img_usage_buf, width=400, height=200))
    else: chart_row.append("")
    
    if img_m_buf or img_usage_buf:
        t_charts = Table([chart_row], colWidths=[420, 420])
        t_charts.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'TOP')]))
        elements.append(t_charts)
        elements.append(Spacer(1, 15))
    
    d_m = [['BULAN', 'SISA BULAN LALU', 'MASUK', 'KELUAR', 'SISA']]
    for r in m_data: d_m.append([r['bln'], f"{r['awal']:,.0f}", f"{r['masuk']:,.0f}", f"{r['keluar']:,.0f}", f"{r['sisa']:,.0f}"])
    if m_data:
        t_masuk = sum(x['masuk'] for x in m_data); t_keluar = sum(x['keluar'] for x in m_data); akhir = m_data[-1]['sisa']
        d_m.append(['TOTAL', '', f"{t_masuk:,.0f}", f"{t_keluar:,.0f}", f"{akhir:,.0f}"])
    t_m = Table(d_m, colWidths=[100, 100, 100, 100, 100])
    rekap_style = [('GRID', (0,0), (-1,-1), 0.5, COLOR_BORDER), ('BACKGROUND', (0,0), (-1,0), COLOR_HEADER_BLUE), ('TEXTCOLOR', (0,0), (-1,0), colors.white), ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'), ('ALIGN', (0,0), (-1,0), 'CENTER'), ('ALIGN', (1,0), (-1,-1), 'RIGHT'), ('FONTSIZE', (0,0), (-1,-1), 9), ('LEFTPADDING', (0,0), (-1,-1), 6), ('RIGHTPADDING', (0,0), (-1,-1), 6)]
    for i in range(1, len(d_m)): bg = COLOR_ROW_EVEN if i % 2 == 0 else COLOR_ROW_ODD; rekap_style.append(('BACKGROUND', (0, i), (-1, i), bg))
    if m_data: rekap_style.append(('BACKGROUND', (0, -1), (-1, -1), COLOR_TOTAL_YELLOW)); rekap_style.append(('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'))
    t_m.setStyle(TableStyle(rekap_style)); elements.append(t_m)
    doc.build(elements); buffer.seek(0)
    return buffer

def generate_excel_one_sheet(conn, lokasi_id, nama_lokasi, start_date_global, end_date_global, excluded_list):
    output = io.BytesIO(); wb = Workbook(); wb.remove(wb.active)
    thin = Border(left=Side('thin'), right=Side('thin'), top=Side('thin'), bottom=Side('thin'))
    date_ranges = split_date_range_by_month(start_date_global, end_date_global)
    
    for idx, (start_date, end_date) in enumerate(date_ranges):
        sheet_name = get_bulan_indonesia(start_date.month)[:3] + f" {start_date.year}"
        ws = wb.create_sheet(sheet_name)
        
        stok_awal = hitung_stok_awal_periode(conn, lokasi_id, start_date)
        df_masuk = pd.read_sql(f"SELECT * FROM bbm_masuk WHERE lokasi_id={lokasi_id} AND tanggal BETWEEN '{start_date}' AND '{end_date}' ORDER BY tanggal", conn)
        df_keluar = pd.read_sql(f"SELECT * FROM bbm_keluar WHERE lokasi_id={lokasi_id} AND tanggal BETWEEN '{start_date}' AND '{end_date}' ORDER BY tanggal", conn)
        if not df_keluar.empty and 'kategori' not in df_keluar.columns: df_keluar['kategori'] = df_keluar['nama_alat'].apply(cek_kategori)
        df_keluar_rpt = filter_non_consumption(df_keluar)
        tm = float(df_masuk['jumlah_liter'].sum()) if not df_masuk.empty else 0.0
        tk_real = float(df_keluar['jumlah_liter'].sum()) if not df_keluar.empty else 0.0
        tk_rpt = float(df_keluar_rpt['jumlah_liter'].sum()) if not df_keluar_rpt.empty else 0.0
        sisa_akhir = stok_awal + tm - tk_real
        
        # REVERT TO 145 FOR EXCEL
        SPLIT_IDX = 145
        df_left = df_keluar_rpt.iloc[:SPLIT_IDX]
        df_right_top = df_keluar_rpt.iloc[SPLIT_IDX:]
        
        ws.merge_cells('A1:N1'); ws['A1'] = f"LAPORAN BBM: {nama_lokasi}"; ws['A1'].font = Font(bold=True, size=14); ws['A1'].alignment = Alignment(horizontal='center')
        ws.merge_cells('A2:N2'); ws['A2'] = f"PERIODE {get_bulan_indonesia(start_date.month)} {start_date.year}"; ws['A2'].font = Font(size=12, bold=True); ws['A2'].alignment = Alignment(horizontal='center')
        
        ws.merge_cells('A4:F4'); ws['A4'] = "PENGGUNAAN BBM (KELUAR)"; ws['A4'].font = Font(bold=True)
        ws['A5'] = "NO"; ws['B5'] = "TGL"; ws['C5'] = "ALAT"; ws['D5'] = "UNIT"; ws['E5'] = "LTR"; ws['F5'] = "KET"
        for c in ['A','B','C','D','E','F']: ws[f'{c}5'].fill = PatternFill("solid", fgColor="2F5496"); ws[f'{c}5'].font = Font(color="FFFFFF", bold=True); ws[f'{c}5'].alignment = Alignment(horizontal='center')
        
        ws.column_dimensions['A'].width = 5; ws.column_dimensions['B'].width = 12; ws.column_dimensions['C'].width = 25
        ws.column_dimensions['D'].width = 15; ws.column_dimensions['E'].width = 10; ws.column_dimensions['F'].width = 30

        current_left_row = 6
        last_date_l = None; is_grey_l = False
        if not df_left.empty:
            for i, r in df_left.iterrows():
                curr_date = r['tanggal']
                if last_date_l is not None and curr_date != last_date_l: is_grey_l = not is_grey_l
                last_date_l = curr_date
                
                ws.cell(current_left_row, 1, i+1); ws.cell(current_left_row, 2, r['tanggal'].strftime('%d/%m'))
                ws.cell(current_left_row, 3, r['nama_alat']); ws.cell(current_left_row, 4, r['no_unit'])
                ws.cell(current_left_row, 5, r['jumlah_liter']); ws.cell(current_left_row, 6, r['keterangan'])
                
                fill_color = PatternFill("solid", fgColor="F2F2F2") if is_grey_l else None
                for cx in range(1,7): cell = ws.cell(current_left_row, cx); cell.border = thin; cell.alignment = Alignment(wrap_text=True, vertical='center'); 
                if fill_color: 
                    for cx in range(1,7): ws.cell(current_left_row, cx).fill = fill_color
                current_left_row += 1
        
        if df_right_top.empty:
            ws.cell(current_left_row, 3, "TOTAL").font=Font(bold=True); 
            ws.cell(current_left_row, 5, tk_rpt).font=Font(bold=True); 
            ws.cell(current_left_row, 5).fill=PatternFill("solid", fgColor="FFFF00")
            for cx in range(1,7): ws.cell(current_left_row, cx).border = thin

        col_start = 9 
        ws.column_dimensions['I'].width = 5; ws.column_dimensions['J'].width = 12; ws.column_dimensions['K'].width = 25
        ws.column_dimensions['L'].width = 15; ws.column_dimensions['M'].width = 10; ws.column_dimensions['N'].width = 30
        current_right_row = 4
        
        if not df_right_top.empty:
            ws.merge_cells(start_row=current_right_row, start_column=col_start, end_row=current_right_row, end_column=col_start+5)
            ws.cell(current_right_row, col_start, "PENGGUNAAN BBM (LANJUTAN)").font = Font(bold=True)
            current_right_row += 1
            headers_r = ["NO", "TGL", "ALAT", "UNIT", "LTR", "KET"]
            for k, h in enumerate(headers_r):
                c = ws.cell(current_right_row, col_start+k, h)
                c.fill = PatternFill("solid", fgColor="2F5496"); c.font = Font(color="FFFFFF", bold=True); c.alignment = Alignment(horizontal='center')
            current_right_row += 1
            
            last_date_r = None; is_grey_r = False
            for idx, (idx_orig, r) in enumerate(df_right_top.iterrows()):
                curr_date = r['tanggal']; last_date_r = curr_date
                if last_date_r is not None and curr_date != last_date_r: is_grey_r = not is_grey_r
                last_date_r = curr_date
                
                ws.cell(current_right_row, col_start, SPLIT_IDX + idx + 1)
                ws.cell(current_right_row, col_start+1, r['tanggal'].strftime('%d/%m'))
                ws.cell(current_right_row, col_start+2, r['nama_alat'])
                ws.cell(current_right_row, col_start+3, r['no_unit'])
                ws.cell(current_right_row, col_start+4, r['jumlah_liter'])
                ws.cell(current_right_row, col_start+5, r['keterangan'])
                
                fill_color = PatternFill("solid", fgColor="F2F2F2") if is_grey_r else None
                for cx in range(6): cell = ws.cell(current_right_row, col_start+cx); cell.border = thin; cell.alignment = Alignment(wrap_text=True, vertical='center'); 
                if fill_color: 
                    for cx in range(6): ws.cell(current_right_row, col_start+cx).fill = fill_color
                current_right_row += 1
            
            ws.cell(current_right_row, col_start+2, "TOTAL").font=Font(bold=True)
            ws.cell(current_right_row, col_start+4, tk_rpt).font=Font(bold=True)
            ws.cell(current_right_row, col_start+4).fill=PatternFill("solid", fgColor="FFFF00")
            for cx in range(6): ws.cell(current_right_row, col_start+cx).border = thin
            current_right_row += 2 
            
        ws.cell(current_right_row, col_start, "BBM MASUK").font = Font(bold=True)
        current_right_row += 1
        headers_m = ['NO', 'TGL', 'SUMBER', 'JNS', 'LTR']
        for k, h in enumerate(headers_m): 
            c = ws.cell(current_right_row, col_start+k, h)
            c.fill = PatternFill("solid", fgColor="2F5496"); c.font = Font(color="FFFFFF", bold=True); c.border = thin
        current_right_row += 1
        
        if not df_masuk.empty:
            for i, r in df_masuk.iterrows():
                ws.cell(current_right_row, col_start, i+1); ws.cell(current_right_row, col_start+1, r['tanggal'].strftime('%d/%m'))
                ws.cell(current_right_row, col_start+2, r['sumber']); ws.cell(current_right_row, col_start+3, r['jenis_bbm']); ws.cell(current_right_row, col_start+4, r['jumlah_liter'])
                for cx in range(5): cell = ws.cell(current_right_row, col_start+cx); cell.border = thin; cell.alignment = Alignment(wrap_text=True)
                current_right_row += 1
        else: ws.cell(current_right_row, col_start+2, "TIDAK ADA DATA"); current_right_row +=1

        ws.cell(current_right_row, col_start+2, "TOTAL").font=Font(bold=True); ws.cell(current_right_row, col_start+4, tm).font=Font(bold=True); ws.cell(current_right_row, col_start+4).fill = PatternFill("solid", fgColor="FFFF00")
        for cx in range(5): ws.cell(current_right_row, col_start+cx).border = thin
        current_right_row += 2
        
        ws.cell(current_right_row, col_start, "RINCIAN PENGGUNAAN BBM").font = Font(bold=True); current_right_row += 1
        df_alat_g, df_truck_g, df_lain_g = segregate_data(df_keluar_rpt, excluded_list)
        
        def write_detail_one_sheet(ws, row, col, title, df, color):
            ws.merge_cells(start_row=row, start_column=col, end_row=row, end_column=col+3)
            c=ws.cell(row, col, title); c.fill=color; c.font=Font(bold=True); c.alignment=Alignment(horizontal='center'); c.border=thin; row+=1
            if not df.empty and 'jumlah_liter' in df.columns:
                grp = df.groupby(['nama_alat', 'no_unit'])['jumlah_liter'].sum().reset_index()
                for _, x in grp.iterrows():
                    ws.merge_cells(start_row=row, start_column=col, end_row=row, end_column=col+2)
                    c1=ws.cell(row, col, f"{x['nama_alat']} {x['no_unit']}"); c1.border=thin; c2=ws.cell(row, col+3, float(x['jumlah_liter'])); c2.border=thin; row+=1
            total_val = float(df['jumlah_liter'].sum()) if not df.empty and 'jumlah_liter' in df.columns else 0
            ws.merge_cells(start_row=row, start_column=col, end_row=row, end_column=col+2)
            c_tot=ws.cell(row, col, "TOTAL"); c_tot.fill=PatternFill("solid", fgColor="FFFF00"); c_tot.border=thin; c_tot.font=Font(bold=True)
            c_val=ws.cell(row, col+3, total_val); c_val.fill=PatternFill("solid", fgColor="FFFF00"); c_val.border=thin; c_val.font=Font(bold=True)
            return row+2

        current_right_row = write_detail_one_sheet(ws, current_right_row, col_start, "TOTAL PENGGUNAAN ALAT BERAT", df_alat_g, PatternFill("solid", fgColor="F4B084"))
        current_right_row = write_detail_one_sheet(ws, current_right_row, col_start, "TOTAL PENGGUNAAN MOBIL & TRUCK", df_truck_g, PatternFill("solid", fgColor="9BC2E6"))
        if not df_lain_g.empty: current_right_row = write_detail_one_sheet(ws, current_right_row, col_start, "TOTAL PENGGUNAAN BBM LAINNYA", df_lain_g, PatternFill("solid", fgColor="FFB6C1"))
        
        ws.merge_cells(start_row=current_right_row, start_column=col_start, end_row=current_right_row, end_column=col_start+2)
        ws.cell(current_right_row, col_start, "RINCIAN SISA STOK BBM").font = Font(bold=True); current_right_row += 1
        
        def write_rekap_row(title, val, color=None):
            nonlocal current_right_row
            ws.merge_cells(start_row=current_right_row, start_column=col_start, end_row=current_right_row, end_column=col_start+2)
            ws.cell(current_right_row, col_start, title).border = thin
            c = ws.cell(current_right_row, col_start+3, val); c.border = thin
            if color: c.fill = PatternFill("solid", fgColor=color)
            current_right_row += 1

        write_rekap_row("SISA BULAN LALU", stok_awal)
        write_rekap_row("TOTAL MASUK", tm)
        write_rekap_row("TOTAL KELUAR", tk_real)
        write_rekap_row("SISA AKHIR", sisa_akhir, "00FF00")
        
        current_right_row += 1
        
        img_buf = generate_chart_for_report(df_alat_g, df_truck_g, width_inch=4.5, height_inch=3.0)
        if img_buf: 
            img = XLImage(img_buf); img.width = 450; img.height = 450
            ws.add_image(img, f'I{current_right_row}')

    ws2 = wb.create_sheet("Rekap Tahunan"); ws2['A1'] = "LAPORAN BBM PERBULAN"; ws2['A1'].font = Font(bold=True, size=14)
    ws2.column_dimensions['A'].width = 25; ws2.column_dimensions['B'].width = 20; ws2.column_dimensions['C'].width = 20; ws2.column_dimensions['D'].width = 20; ws2.column_dimensions['E'].width = 20
    m_data = []
    stok_run = hitung_stok_awal_periode(conn, lokasi_id, start_date_global)
    curr = start_date_global.replace(day=1); end_limit = end_date_global.replace(day=1)
    while curr <= end_limit:
        m = curr.month; y = curr.year
        q_in = f"SELECT SUM(jumlah_liter) FROM bbm_masuk WHERE lokasi_id={lokasi_id} AND MONTH(tanggal)={m} AND YEAR(tanggal)={y}"
        q_out = f"SELECT SUM(jumlah_liter) FROM bbm_keluar WHERE lokasi_id={lokasi_id} AND MONTH(tanggal)={m} AND YEAR(tanggal)={y}"
        cursor = conn.cursor()
        cursor.execute(q_in); res_in = cursor.fetchone(); mi = float(res_in[0]) if res_in and res_in[0] else 0.0
        cursor.execute(q_out); res_out = cursor.fetchone(); mo = float(res_out[0]) if res_out and res_out[0] else 0.0
        prev = stok_run; stok_run = prev + mi - mo
        m_data.append({'bln': curr.strftime("%B %Y"), 'awal': prev, 'masuk': mi, 'keluar': mo, 'sisa': stok_run, 'bulan_nama': get_bulan_indonesia(m)[:3]})
        curr = curr + relativedelta(months=1)
    df_m = pd.DataFrame(m_data)
    img_m_buf = generate_monthly_chart(df_m)
    if img_m_buf: img2 = XLImage(img_m_buf); img2.width=500; img2.height=250; ws2.add_image(img2, 'A3')
    r2 = 18; headers = ['BULAN', 'SISA BULAN LALU', 'MASUK', 'KELUAR', 'SISA']
    for i, h in enumerate(headers): c=ws2.cell(r2, i+1, h); c.border=thin; c.fill=PatternFill("solid", fgColor="D3D3D3")
    r2+=1
    for r in m_data:
        vals = [r['bln'], r['awal'], r['masuk'], r['keluar'], r['sisa']]
        for i, v in enumerate(vals): c=ws2.cell(r2, i+1, v); c.border=thin
        r2+=1
    if m_data:
        t_masuk = sum(x['masuk'] for x in m_data); t_keluar = sum(x['keluar'] for x in m_data); akhir = m_data[-1]['sisa']
        ws2.cell(r2, 1, "TOTAL").font = Font(bold=True); ws2.cell(r2, 3, t_masuk).font = Font(bold=True); ws2.cell(r2, 4, t_keluar).font = Font(bold=True); ws2.cell(r2, 5, akhir).font = Font(bold=True)
        for i in range(1, 6): c = ws2.cell(r2, i); c.fill = PatternFill("solid", fgColor="FFD966"); c.border = thin

    # --- ADD CHART TO EXCEL ONE SHEET (RIGHT OF MONTHLY CHART) ---
    df_keluar_all = pd.read_sql(f"SELECT * FROM bbm_keluar WHERE lokasi_id={lokasi_id} AND tanggal BETWEEN '{start_date_global}' AND '{end_date_global}'", conn)
    if not df_keluar_all.empty:
        if 'kategori' not in df_keluar_all.columns: df_keluar_all['kategori'] = df_keluar_all['nama_alat'].apply(cek_kategori)
        df_keluar_rpt = filter_non_consumption(df_keluar_all)
        df_alat_t, df_truck_t, _ = segregate_data(df_keluar_rpt, excluded_list)
        img_usage = generate_chart_for_report(df_alat_t, df_truck_t, width_inch=7, height_inch=3.5)
        if img_usage:
            img3 = XLImage(img_usage); img3.width=500; img3.height=250
            ws2.add_image(img3, 'H3')

    wb.save(output); output.seek(0)
    return output

def generate_docx_one_sheet(conn, lokasi_id, nama_lokasi, start_date_global, end_date_global, excluded_list):
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(55.88) 
    section.page_width = Inches(14) 
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    
    date_ranges = split_date_range_by_month(start_date_global, end_date_global)
    
    for idx, (start_date, end_date) in enumerate(date_ranges):
        if idx > 0: doc.add_page_break()
        p = doc.add_paragraph(f"LAPORAN BBM: {nama_lokasi}"); 
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold=True; p.runs[0].font.size=Pt(14)
        p.paragraph_format.space_after = Pt(0)
        p2 = doc.add_paragraph(f"PERIODE {get_bulan_indonesia(start_date.month)} {start_date.year}")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.runs[0].bold=True; p2.runs[0].font.size=Pt(12)
        p2.paragraph_format.space_after = Pt(6) # Sedikit jarak setelah judul periode
        
        stok_awal = hitung_stok_awal_periode(conn, lokasi_id, start_date)
        df_keluar = pd.read_sql(f"SELECT * FROM bbm_keluar WHERE lokasi_id={lokasi_id} AND tanggal BETWEEN '{start_date}' AND '{end_date}' ORDER BY tanggal", conn)
        df_masuk = pd.read_sql(f"SELECT * FROM bbm_masuk WHERE lokasi_id={lokasi_id} AND tanggal BETWEEN '{start_date}' AND '{end_date}' ORDER BY tanggal", conn)
        if not df_keluar.empty and 'kategori' not in df_keluar.columns: df_keluar['kategori'] = df_keluar['nama_alat'].apply(cek_kategori)
        df_keluar_rpt = filter_non_consumption(df_keluar)
        df_alat_g, df_truck_g, df_lain_g = segregate_data(df_keluar_rpt, excluded_list)
        tm = float(df_masuk['jumlah_liter'].sum()) if not df_masuk.empty else 0.0
        tk_real = float(df_keluar['jumlah_liter'].sum()) if not df_keluar.empty else 0.0
        tk_rpt = float(df_keluar_rpt['jumlah_liter'].sum()) if not df_keluar_rpt.empty else 0.0
        sisa_akhir = stok_awal + tm - tk_real
        
        # REVERT TO 145 FOR DOCX
        SPLIT_IDX = 145
        df_left = df_keluar_rpt.iloc[:SPLIT_IDX]
        df_right_top = df_keluar_rpt.iloc[SPLIT_IDX:]
        
        layout_table = doc.add_table(rows=1, cols=2)
        layout_table.autofit = False
        layout_table.columns[0].width = Inches(6.5)
        layout_table.columns[1].width = Inches(6.5)
        cell_left = layout_table.cell(0, 0); cell_right = layout_table.cell(0, 1)
        
        p_judul_kiri = cell_left.add_paragraph(f"PENGGUNAAN BBM (1-{SPLIT_IDX})")
        p_judul_kiri.paragraph_format.space_after = Pt(2)
        t_left = cell_left.add_table(rows=1, cols=5)
        t_left.style = 'Table Grid'
        hdr = t_left.rows[0].cells; hdr[0].text="NO"; hdr[1].text="TGL"; hdr[2].text="ALAT"; hdr[3].text="UNIT"; hdr[4].text="LTR"
        for c in hdr: 
             set_cell_bg(c, "2F5496"); p = c.paragraphs[0]; p.runs[0].font.size = Pt(7)
             p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = Pt(8)
        
        last_date_l = None; is_grey_l = False
        if not df_left.empty:
            for i, r in df_left.iterrows():
                curr_date = r['tanggal']
                if last_date_l is not None and curr_date != last_date_l: is_grey_l = not is_grey_l
                last_date_l = curr_date
                row = t_left.add_row().cells
                row[0].text = str(i+1); row[1].text = r['tanggal'].strftime('%d/%m'); row[2].text = r['nama_alat']; row[3].text = r['no_unit']; row[4].text = str(r['jumlah_liter'])
                for c in row: 
                    p = c.paragraphs[0]; p.runs[0].font.size = Pt(7); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = Pt(8)
                    if is_grey_l: set_cell_bg(c, "F2F2F2")
        
        if df_right_top.empty:
             row = t_left.add_row().cells; row[2].text = "TOTAL"; row[4].text = f"{tk_rpt:.0f}"
             set_cell_bg(row[4], "FFFF00"); row[4].paragraphs[0].paragraph_format.space_after = Pt(0)

        if not df_right_top.empty:
            p_judul_kanan = cell_right.add_paragraph("PENGGUNAAN BBM (LANJUTAN)")
            p_judul_kanan.paragraph_format.space_after = Pt(2)
            t_rt = cell_right.add_table(rows=1, cols=5)
            t_rt.style = 'Table Grid'
            hdr = t_rt.rows[0].cells; hdr[0].text="NO"; hdr[1].text="TGL"; hdr[2].text="ALAT"; hdr[3].text="UNIT"; hdr[4].text="LTR"
            for c in hdr: 
                set_cell_bg(c, "2F5496"); p = c.paragraphs[0]; p.runs[0].font.size = Pt(7)
                p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = Pt(8)
            
            last_date_r = None; is_grey_r = False
            for idx, (idx_orig, r) in enumerate(df_right_top.iterrows()):
                curr_date = r['tanggal']
                if last_date_r is not None and curr_date != last_date_r: is_grey_r = not is_grey_r
                last_date_r = curr_date
                row = t_rt.add_row().cells
                num = SPLIT_IDX + idx + 1
                row[0].text = str(num); row[1].text = r['tanggal'].strftime('%d/%m'); row[2].text = r['nama_alat']; row[3].text = r['no_unit']; row[4].text = str(r['jumlah_liter'])
                for c in row: 
                     p = c.paragraphs[0]; p.runs[0].font.size = Pt(7); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = Pt(8)
                     if is_grey_r: set_cell_bg(c, "F2F2F2")
            
            row = t_rt.add_row().cells; row[2].text = "TOTAL"; row[4].text = f"{tk_rpt:.0f}"
            set_cell_bg(row[4], "FFFF00"); row[4].paragraphs[0].paragraph_format.space_after = Pt(0)
            
        p_masuk = cell_right.add_paragraph("BBM MASUK")
        p_masuk.paragraph_format.space_before = Pt(6) # Jarak kecil sebelum judul
        p_masuk.paragraph_format.space_after = Pt(2)

        t_m = cell_right.add_table(rows=1, cols=4); t_m.style = 'Table Grid'
        hdr = t_m.rows[0].cells; hdr[0].text="TGL"; hdr[1].text="SUMBER"; hdr[2].text="JNS"; hdr[3].text="LTR"
        for c in hdr: set_cell_bg(c, "2F5496"); c.paragraphs[0].runs[0].font.size = Pt(7); c.paragraphs[0].paragraph_format.space_after=Pt(0)
        if not df_masuk.empty:
            for i, r in df_masuk.iterrows():
                row = t_m.add_row().cells; row[0].text=r['tanggal'].strftime('%d/%m'); row[1].text=r['sumber']; row[2].text=r['jenis_bbm']; row[3].text=str(r['jumlah_liter'])
                for c in row: c.paragraphs[0].runs[0].font.size = Pt(7); c.paragraphs[0].paragraph_format.space_after=Pt(0)
        
        row = t_m.add_row().cells; row[1].text = "TOTAL"; row[3].text = f"{tm:.0f}"; set_cell_bg(row[3], "FFFF00"); row[3].paragraphs[0].paragraph_format.space_after=Pt(0)
        
        p_rincian = cell_right.add_paragraph("RINCIAN PENGGUNAAN BBM")
        p_rincian.paragraph_format.space_before = Pt(6)
        p_rincian.paragraph_format.space_after = Pt(2)

        def add_docx_detail_rekap(title, df, color):
            p = cell_right.add_paragraph(title); p.runs[0].bold=True; p.paragraph_format.space_after=Pt(0)
            t = cell_right.add_table(rows=1, cols=2); t.style='Table Grid'
            if not df.empty:
                grp = df.groupby(['nama_alat', 'no_unit'])['jumlah_liter'].sum().reset_index()
                for _, r in grp.iterrows():
                    row = t.add_row().cells; row[0].text = f"{r['nama_alat']} {r['no_unit']}"; row[1].text = f"{r['jumlah_liter']:.0f}"
                    for c in row: c.paragraphs[0].runs[0].font.size = Pt(7); c.paragraphs[0].paragraph_format.space_after=Pt(0)
            row = t.add_row().cells; row[0].text="TOTAL"; row[1].text=f"{df['jumlah_liter'].sum():.0f}"
            set_cell_bg(row[0], color); set_cell_bg(row[1], "FFFF00"); row[0].paragraphs[0].paragraph_format.space_after=Pt(0)
            # Hilangkan baris kosong di sini

        add_docx_detail_rekap("TOTAL ALAT BERAT", df_alat_g, "F4B084")
        add_docx_detail_rekap("TOTAL MOBIL & TRUCK", df_truck_g, "9BC2E6")
        if not df_lain_g.empty: add_docx_detail_rekap("LAINNYA", df_lain_g, "FFB6C1")

        p_sisa = cell_right.add_paragraph("RINCIAN SISA STOK BBM")
        p_sisa.paragraph_format.space_before = Pt(6)
        p_sisa.paragraph_format.space_after = Pt(2)

        t_s = cell_right.add_table(rows=4, cols=2); t_s.style = 'Table Grid'
        t_s.cell(0,0).text = "SISA BULAN LALU"; t_s.cell(0,1).text = f"{stok_awal:.0f}"
        t_s.cell(1,0).text = "TOTAL MASUK"; t_s.cell(1,1).text = f"{tm:.0f}"
        t_s.cell(2,0).text = "TOTAL KELUAR"; t_s.cell(2,1).text = f"{tk_real:.0f}"
        t_s.cell(3,0).text = "SISA AKHIR"; t_s.cell(3,1).text = f"{sisa_akhir:.0f}"; set_cell_bg(t_s.cell(3,1), "00FF00")
        for r in t_s.rows: 
             for c in r.cells: c.paragraphs[0].paragraph_format.space_after=Pt(0); c.paragraphs[0].runs[0].font.size=Pt(7)
        
        # --- CHART IN DOCX ONE SHEET (DAILY PART) ---
        img_buf = generate_chart_for_report(df_alat_g, df_truck_g, width_inch=3.5, height_inch=2.5)
        if img_buf: 
            cell_right.add_paragraph("")
            cell_right.add_paragraph().add_run().add_picture(img_buf, width=Cm(8))

    doc.add_page_break(); p_title = doc.add_paragraph("LAPORAN BBM PERBULAN"); p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_title.runs[0].bold=True; p_title.runs[0].font.size=Pt(14)
    m_data = []; stok_run = hitung_stok_awal_periode(conn, lokasi_id, start_date_global)
    curr = start_date_global.replace(day=1); end_limit = end_date_global.replace(day=1)
    while curr <= end_limit:
        m = curr.month; y = curr.year
        q_in = f"SELECT SUM(jumlah_liter) FROM bbm_masuk WHERE lokasi_id={lokasi_id} AND MONTH(tanggal)={m} AND YEAR(tanggal)={y}"
        q_out = f"SELECT SUM(jumlah_liter) FROM bbm_keluar WHERE lokasi_id={lokasi_id} AND MONTH(tanggal)={m} AND YEAR(tanggal)={y}"
        cursor = conn.cursor()
        cursor.execute(q_in); res_in = cursor.fetchone(); mi = float(res_in[0]) if res_in and res_in[0] else 0.0
        cursor.execute(q_out); res_out = cursor.fetchone(); mo = float(res_out[0]) if res_out and res_out[0] else 0.0
        prev = stok_run; stok_run = prev + mi - mo
        m_data.append({'bln': curr.strftime("%B %Y"), 'awal': prev, 'masuk': mi, 'keluar': mo, 'sisa': stok_run, 'bulan_nama': get_bulan_indonesia(m)[:3]})
        curr = curr + relativedelta(months=1)

    df_m = pd.DataFrame(m_data)
    if not df_m.empty:
        img_m_buf = generate_monthly_chart(df_m)
        if img_m_buf: doc.add_paragraph().add_run().add_picture(img_m_buf, width=Cm(16))
    doc.add_paragraph("RINCIAN MASUK DAN PENGGUNAAN SOLAR PERBULANNYA", style='Heading 4')
    tbl_month = doc.add_table(rows=1, cols=5); tbl_month.style='Table Grid'
    h_month = tbl_month.rows[0].cells
    for i, t in enumerate(['BULAN', 'SISA BULAN LALU', 'MASUK', 'KELUAR', 'SISA']): h_month[i].text=t; set_cell_bg(h_month[i], "2F5496")
    for r in m_data:
        row = tbl_month.add_row().cells; row[0].text=r['bln']; row[1].text=f"{r['awal']:.0f}"; row[2].text=f"{r['masuk']:.0f}"; row[3].text=f"{r['keluar']:.0f}"; row[4].text=f"{r['sisa']:.0f}"
    
    if m_data:
        t_masuk = sum(x['masuk'] for x in m_data); t_keluar = sum(x['keluar'] for x in m_data); akhir = m_data[-1]['sisa']
        row = tbl_month.add_row().cells; row[0].text = "TOTAL"; row[2].text = f"{t_masuk:,.0f}"; row[3].text = f"{t_keluar:,.0f}"; row[4].text = f"{akhir:,.0f}"
        for c in row:
            set_cell_bg(c, "FFD966")
            if len(c.paragraphs)>0 and len(c.paragraphs[0].runs)>0: c.paragraphs[0].runs[0].bold = True
            elif len(c.paragraphs)>0: c.paragraphs[0].add_run(c.text).font.bold = True

    # --- ADD CHART TO DOCX ONE SHEET (BELOW TABLE IN SUMMARY) ---
    df_keluar_all = pd.read_sql(f"SELECT * FROM bbm_keluar WHERE lokasi_id={lokasi_id} AND tanggal BETWEEN '{start_date_global}' AND '{end_date_global}'", conn)
    if not df_keluar_all.empty:
        if 'kategori' not in df_keluar_all.columns: df_keluar_all['kategori'] = df_keluar_all['nama_alat'].apply(cek_kategori)
        df_keluar_rpt = filter_non_consumption(df_keluar_all)
        df_alat_t, df_truck_t, _ = segregate_data(df_keluar_rpt, excluded_list)
        img_usage = generate_chart_for_report(df_alat_t, df_truck_t, width_inch=7, height_inch=3.5)
        if img_usage:
            doc.add_paragraph().add_run().add_picture(img_usage, width=Cm(16))

    buffer = io.BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer

def main():
    try: 
        conn = init_connection(); cursor = conn.cursor(buffered=True) 
        try: cursor.execute("SELECT stok_awal FROM lokasi_proyek LIMIT 1"); cursor.fetchall()
        except: cursor.execute("ALTER TABLE lokasi_proyek ADD COLUMN stok_awal FLOAT DEFAULT 0"); conn.commit()
        try: cursor.execute("SELECT kunci_lokasi FROM lokasi_proyek LIMIT 1"); cursor.fetchall()
        except: cursor.execute("ALTER TABLE lokasi_proyek ADD COLUMN kunci_lokasi VARCHAR(255) DEFAULT '123'"); conn.commit()
        cursor.execute("""CREATE TABLE IF NOT EXISTS log_aktivitas (id INT AUTO_INCREMENT PRIMARY KEY, lokasi_id INT, tanggal DATETIME DEFAULT CURRENT_TIMESTAMP, kategori VARCHAR(50), deskripsi TEXT, affected_ids TEXT)""")
        try: cursor.execute("SELECT affected_ids FROM log_aktivitas LIMIT 1"); cursor.fetchall()
        except: 
             try: cursor.execute("ALTER TABLE log_aktivitas ADD COLUMN affected_ids TEXT"); conn.commit()
             except: pass
        cursor.execute("""CREATE TABLE IF NOT EXISTS rekap_exclude (id INT AUTO_INCREMENT PRIMARY KEY, lokasi_id INT, nama_unit_full VARCHAR(255))"""); conn.commit()
    except Exception as e: st.error(f"Database Error: {e}"); st.stop()

    password_rahasia = "123" 
    if "logged_in" not in st.session_state: st.session_state.logged_in = False
    if "active_project_id" not in st.session_state: st.session_state.active_project_id = None
    if "active_project_name" not in st.session_state: st.session_state.active_project_name = None

    if not st.session_state.logged_in:
        st.markdown("<h1 style='text-align: center;'>🔐 Login Admin</h1>", unsafe_allow_html=True)
        c1, c2, c3 = st.columns([1,2,1])
        with c2:
            pwd = st.text_input("Masukkan Password Admin:", type="password")
            if st.button("Login", use_container_width=True):
                if pwd == password_rahasia: st.session_state.logged_in = True; st.rerun()
                else: st.error("Password Salah!")
        st.stop() 

    if st.session_state.active_project_id is None:
        st.title("🗂️ Menu Utama")
        st.write("Selamat Datang, Admin. Silakan pilih lokasi proyek atau buat baru.")
        st.divider()
        col_left, col_right = st.columns(2, gap="large")
        with col_left:
            with st.container(border=True):
                st.subheader("📂 Masuk ke Lokasi Proyek")
                try: df_lokasi = pd.read_sql("SELECT * FROM lokasi_proyek", conn)
                except: df_lokasi = pd.DataFrame()
                if not df_lokasi.empty:
                    pilih_nama = st.selectbox("Pilih Lokasi:", df_lokasi['nama_tempat'])
                    input_pass_lokasi = st.text_input("Password Lokasi:", type="password", key="pass_enter")
                    if st.button("Masuk Lokasi", type="primary", use_container_width=True):
                        data_lok = df_lokasi[df_lokasi['nama_tempat'] == pilih_nama].iloc[0]
                        if input_pass_lokasi == data_lok['kunci_lokasi']:
                            st.session_state.active_project_id = int(data_lok['id']); st.session_state.active_project_name = data_lok['nama_tempat']; st.success(f"Berhasil masuk ke {pilih_nama}"); st.rerun()
                        else: st.error("Password Lokasi Salah!")
                else: st.info("Belum ada data lokasi.")
        with col_right:
            with st.container(border=True):
                st.subheader("➕ Buat Lokasi Baru")
                new_name = st.text_input("Nama Lokasi Baru")
                new_pass = st.text_input("Buat Password Lokasi", type="password", key="pass_create")
                st.warning("⚠️ **Password nya diingat baik-baik !**")
                if st.button("Simpan Lokasi Baru", use_container_width=True):
                    if new_name and new_pass:
                        try:
                            cursor.execute("INSERT INTO lokasi_proyek (nama_tempat, kunci_lokasi) VALUES (%s, %s)", (new_name, new_pass)); conn.commit(); st.success("Lokasi berhasil dibuat!"); st.rerun()
                        except Exception as e: st.error(f"Gagal membuat lokasi: {e}")
                    else: st.error("Nama dan Password wajib diisi!")
        st.stop() 

    lokasi_id = st.session_state.active_project_id; nama_proyek = st.session_state.active_project_name
    cursor.execute("SELECT stok_awal FROM lokasi_proyek WHERE id=%s", (lokasi_id,)); stok_awal_modal = cursor.fetchone()[0]
    
    with st.sidebar:
        st.header(f"📍 {nama_proyek}")
        if st.button("⬅️ Kembali ke Menu Utama", use_container_width=True): st.session_state.active_project_id = None; st.session_state.active_project_name = None; st.rerun()
        st.divider(); st.info("ℹ️ Stok Awal dihitung otomatis berdasarkan history transaksi dan filter tanggal.")

    df_masuk_all = pd.read_sql(f"SELECT * FROM bbm_masuk WHERE lokasi_id={lokasi_id}", conn)
    df_keluar_all = pd.read_sql(f"SELECT * FROM bbm_keluar WHERE lokasi_id={lokasi_id}", conn)
    df_log = pd.read_sql(f"SELECT * FROM log_aktivitas WHERE lokasi_id={lokasi_id}", conn)
    df_ex = pd.read_sql(f"SELECT nama_unit_full FROM rekap_exclude WHERE lokasi_id={lokasi_id}", conn)
    excluded_list = df_ex['nama_unit_full'].tolist() if not df_ex.empty else []

    if not df_masuk_all.empty: df_masuk_all['HARI'] = pd.to_datetime(df_masuk_all['tanggal']).apply(get_hari_indonesia)
    if not df_keluar_all.empty: df_keluar_all['HARI'] = pd.to_datetime(df_keluar_all['tanggal']).apply(get_hari_indonesia)

    st.title(f"🚜 Dashboard: {nama_proyek}")
    t1, t2, t3 = st.tabs(["📝 Input & History", "📊 Laporan & Grafik", "🖨️ Export Dokumen"])
    
    with t1:
        st.subheader("Input Transaksi BBM")
        mode_transaksi = st.radio("Pilih Jenis Transaksi:", ["📥 BBM MASUK", "📤 PENGGUNAAN BBM", "🔄 PINJAM / TRANSFER ANTAR UNIT"], horizontal=True)
        with st.container(border=True):
            if mode_transaksi == "📥 BBM MASUK":
                with st.form("form_masuk"):
                    c1, c2 = st.columns(2)
                    with c1: tg = st.date_input("Tanggal Masuk"); sm = st.text_input("Sumber / Supplier")
                    with c2: jn = st.selectbox("Jenis BBM", ["Dexlite","Solar","Bensin"]); jl = st.number_input("Jumlah Liter", 0.0); kt = st.text_area("Keterangan")
                    if st.form_submit_button("Simpan BBM Masuk"):
                        cursor.execute("SELECT id FROM bbm_masuk WHERE lokasi_id=%s AND tanggal=%s AND sumber=%s AND jumlah_liter=%s", (lokasi_id, tg, sm, jl))
                        if cursor.fetchall(): st.warning("⚠️ Data serupa sudah ada!") 
                        cursor.execute("INSERT INTO bbm_masuk (lokasi_id, tanggal, sumber, jenis_bbm, jumlah_liter, keterangan) VALUES (%s,%s,%s,%s,%s,%s)", (lokasi_id, tg, sm, jn, jl, kt)); conn.commit(); st.success("Data Masuk Tersimpan!"); st.rerun()
            elif mode_transaksi == "📤 PENGGUNAAN BBM":
                with st.form("form_keluar"):
                    c1, c2 = st.columns(2)
                    with c1: tg_p = st.date_input("Tanggal Pakai"); al = st.text_input("Nama Alat/Kendaraan"); un = st.text_input("Kode Unit (Ex: DT-01)")
                    with c2: jl_p = st.number_input("Liter Digunakan", 0.0); kt_p = st.text_area("Keterangan / Lokasi Kerja")
                    if st.form_submit_button("Simpan Penggunaan"):
                        cursor.execute("SELECT id FROM bbm_keluar WHERE lokasi_id=%s AND tanggal=%s AND nama_alat=%s AND no_unit=%s AND jumlah_liter=%s", (lokasi_id, tg_p, al, un, jl_p))
                        if cursor.fetchall(): st.warning("⚠️ Data serupa sudah ada!") 
                        cursor.execute("INSERT INTO bbm_keluar (lokasi_id, tanggal, nama_alat, no_unit, jumlah_liter, keterangan) VALUES (%s,%s,%s,%s,%s,%s)", (lokasi_id, tg_p, al, un, jl_p, kt_p)); conn.commit(); st.success("Data Penggunaan Tersimpan!"); st.rerun()
            elif mode_transaksi == "🔄 PINJAM / TRANSFER ANTAR UNIT":
                st.info("ℹ️ Mode ini memindahkan liter dari satu unit ke unit lain.")
                with st.form("form_transfer"):
                    c1, c2 = st.columns(2)
                    with c1: tgl_tf = st.date_input("Tanggal Transfer"); donor_alat = st.text_input("DARI ALAT (Pemberi/Donor)"); donor_unit = st.text_input("No Unit Donor")
                    with c2: liter_tf = st.number_input("Jumlah Liter Dipinjam", min_value=0.0); recv_alat = st.text_input("KE ALAT (Penerima)"); recv_unit = st.text_input("No Unit Penerima"); ket_tf = st.text_area("Keterangan Tambahan")
                    if st.form_submit_button("Proses Transfer"):
                        if liter_tf > 0 and donor_alat and recv_alat:
                            ket_donor = f"Transfer ke {recv_alat} {recv_unit}. {ket_tf}"; cursor.execute("INSERT INTO bbm_keluar (lokasi_id, tanggal, nama_alat, no_unit, jumlah_liter, keterangan) VALUES (%s,%s,%s,%s,%s,%s)", (lokasi_id, tgl_tf, donor_alat, donor_unit, -liter_tf, ket_donor))
                            ket_recv = f"Pinjam dari {donor_alat} {donor_unit}. {ket_tf}"; cursor.execute("INSERT INTO bbm_keluar (lokasi_id, tanggal, nama_alat, no_unit, jumlah_liter, keterangan) VALUES (%s,%s,%s,%s,%s,%s)", (lokasi_id, tgl_tf, recv_alat, recv_unit, liter_tf, ket_recv)); conn.commit(); st.success(f"Berhasil transfer {liter_tf}L"); st.rerun()
                        else: st.error("Mohon lengkapi nama alat dan jumlah liter!")
        st.divider()
        with st.expander("⏳ RIWAYAT INPUT & UNDO (Filter & Sortir)", expanded=True):
            c_f1, c_f2, c_f3 = st.columns(3)
            filter_tipe = c_f1.multiselect("Filter Jenis:", ["MASUK", "PAKAI", "TRANSFER", "KOREKSI"], default=["MASUK", "PAKAI", "TRANSFER", "KOREKSI"])
            filter_sort = c_f2.selectbox("Urutkan:", ["Waktu Input Terbaru (ID)", "Waktu Input Terlama (ID)", "Tanggal Laporan Terbaru", "Tanggal Laporan Terlama"])
            use_date_filter = c_f3.checkbox("Filter Tanggal Tertentu")
            date_val = c_f3.date_input("Pilih Tanggal", value=datetime.date.today(), disabled=not use_date_filter)
            history_data = []
            if not df_masuk_all.empty:
                temp_m = df_masuk_all.copy(); temp_m['Tipe'] = 'MASUK'; temp_m['Kategori_Filter'] = 'MASUK'; temp_m['Detail'] = temp_m['sumber'] + " (" + temp_m['jenis_bbm'] + ")"; temp_m['Label_History'] = "📥 BBM MASUK (Beli)"
                history_data.append(temp_m[['id', 'tanggal', 'Tipe', 'Detail', 'jumlah_liter', 'Label_History', 'keterangan', 'Kategori_Filter']])
            if not df_keluar_all.empty:
                temp_k = df_keluar_all.copy(); temp_k['Tipe'] = 'KELUAR'; temp_k['Detail'] = temp_k['nama_alat'] + " " + temp_k['no_unit']
                def tentukan_label(row):
                    liter = row['jumlah_liter']; ket = str(row['keterangan']).lower()
                    if liter < 0: return ("🔄 TRANSFER KELUAR (Donor)", "TRANSFER")
                    elif "pinjam" in ket or "transfer" in ket: return ("🔄 TRANSFER MASUK (Terima)", "TRANSFER")
                    else: return ("📤 PENGGUNAAN (Pakai)", "PAKAI")
                applied = temp_k.apply(tentukan_label, axis=1)
                temp_k['Label_History'] = [x[0] for x in applied]; temp_k['Kategori_Filter'] = [x[1] for x in applied]
                history_data.append(temp_k[['id', 'tanggal', 'Tipe', 'Detail', 'jumlah_liter', 'Label_History', 'keterangan', 'Kategori_Filter']])
            if not df_log.empty:
                temp_l = df_log.copy(); temp_l['Tipe'] = 'LOG'; temp_l['Kategori_Filter'] = 'KOREKSI'; temp_l['Detail'] = temp_l['kategori']; temp_l['jumlah_liter'] = 0; temp_l['Label_History'] = "🛠️ ADMIN/KOREKSI"; temp_l['keterangan'] = temp_l['deskripsi']
                temp_l['affected_ids_val'] = temp_l['affected_ids'] if 'affected_ids' in temp_l.columns else None
                history_data.append(temp_l[['id', 'tanggal', 'Tipe', 'Detail', 'jumlah_liter', 'Label_History', 'keterangan', 'Kategori_Filter', 'affected_ids_val']])
            if history_data:
                df_history = pd.concat(history_data)
                if filter_tipe: df_history = df_history[df_history['Kategori_Filter'].isin(filter_tipe)]
                if use_date_filter: df_history = df_history[pd.to_datetime(df_history['tanggal']).dt.date == date_val]
                if "Waktu Input Terbaru" in filter_sort: df_history.sort_values(by='id', ascending=False, inplace=True)
                elif "Waktu Input Terlama" in filter_sort: df_history.sort_values(by='id', ascending=True, inplace=True)
                elif "Tanggal Laporan Terbaru" in filter_sort: df_history.sort_values(by='tanggal', ascending=False, inplace=True)
                elif "Tanggal Laporan Terlama" in filter_sort: df_history.sort_values(by='tanggal', ascending=True, inplace=True)
                limit_view = st.selectbox("Jumlah Data Ditampilkan:", ["10", "50", "100", "SEMUA"])
                df_view = df_history.head(int(limit_view)) if limit_view != "SEMUA" else df_history
                st.write(f"Menampilkan **{len(df_view)}** data."); st.info("💡 **Catatan Undo Transfer:** Jika membatalkan transfer, pastikan Anda menghapus **KEDUA** baris (Baris 'Donor' dan Baris 'Terima') agar stok kembali seimbang.")
                for index, row in df_view.iterrows():
                    col_a, col_b, col_c, col_d, col_e = st.columns([2, 2, 3, 1, 1])
                    with col_a: 
                        lbl = row['Label_History']
                        if "MASUK" in lbl and "TRANSFER" not in lbl: st.success(lbl)
                        elif "PENGGUNAAN" in lbl: st.warning(lbl)
                        elif "TRANSFER" in lbl: st.info(lbl)
                        elif "ADMIN" in lbl: st.error(lbl)
                        else: st.write(lbl)
                    with col_b: st.write(f"📅 {row['tanggal'].strftime('%d/%m/%Y')}")
                    with col_c: st.write(f"**{row['Detail']}**"); st.caption(f"ket: {row['keterangan']}")
                    with col_d: st.write(f"{row['jumlah_liter']:,.0f} L") if row['Tipe'] != 'LOG' else st.write("-")
                    with col_e:
                        if row['Tipe'] != 'LOG':
                            if st.button("❌ Hapus", key=f"hist_del_{row['Tipe']}_{row['id']}"):
                                table_del = "bbm_masuk" if row['Tipe'] == 'MASUK' else "bbm_keluar"
                                cursor.execute(f"DELETE FROM {table_del} WHERE id=%s", (row['id'],)); conn.commit(); st.success("Data berhasil dihapus!"); st.rerun()
                        else:
                            if st.button("↩️ Undo", key=f"hist_undo_{row['id']}"):
                                try:
                                    matches = re.findall(r"'(.*?)'", row['keterangan'])
                                    if len(matches) >= 2:
                                        old_val = matches[0]; affected_ids = row.get('affected_ids_val')
                                        if affected_ids:
                                            field = "nama_alat" if "NAMA ALAT" in row['Detail'] else "no_unit"
                                            cursor.execute(f"UPDATE bbm_keluar SET {field}='{old_val}' WHERE id IN ({affected_ids})")
                                    cursor.execute("DELETE FROM log_aktivitas WHERE id=%s", (row['id'],)); conn.commit(); st.success(f"Berhasil Undo."); st.rerun()
                                except Exception as e: st.error(f"Error Undo: {e}")
                    st.markdown("---")
            else: st.write("Belum ada riwayat input.")

    with t2:
        st.markdown("### 🗓️ Filter Periode Laporan")
        c_p1, c_p2 = st.columns(2)
        today = datetime.date.today(); first_day = today.replace(day=1)
        if 't2_start' not in st.session_state: st.session_state.t2_start = first_day
        if 't2_end' not in st.session_state: st.session_state.t2_end = today
        with c_p1: start_rep = st.date_input("Mulai Tanggal", value=st.session_state.t2_start, key="pick_start_t2"); st.session_state.t2_start = start_rep
        with c_p2: end_rep = st.date_input("Sampai Tanggal", value=st.session_state.t2_end, key="pick_end_t2"); st.session_state.t2_end = end_rep
        
        df_masuk_rep = df_masuk_all[(pd.to_datetime(df_masuk_all['tanggal']).dt.date >= start_rep) & (pd.to_datetime(df_masuk_all['tanggal']).dt.date <= end_rep)]
        df_keluar_rep = df_keluar_all[(pd.to_datetime(df_keluar_all['tanggal']).dt.date >= start_rep) & (pd.to_datetime(df_keluar_all['tanggal']).dt.date <= end_rep)]
        stok_awal_periode_val = hitung_stok_awal_periode(conn, lokasi_id, start_rep)
        tm_rep = float(df_masuk_rep['jumlah_liter'].sum()); tk_rep = float(df_keluar_rep['jumlah_liter'].sum())
        sisa_rep = stok_awal_periode_val + tm_rep - tk_rep

        st.markdown(f"""<div style="background-color:#d4edda;padding:15px;border-radius:10px;border:1px solid #c3e6cb;text-align:center;margin-bottom:20px;margin-top:10px;"><h2 style="color:#155724;margin:0;">💰 SISA STOK PERIODE INI: {sisa_rep:,.2f} Liter</h2><span style="color:#155724;font-weight:bold;">(Sisa Bulan Lalu: {stok_awal_periode_val:,.0f} + Masuk: {tm_rep:,.0f} - Keluar: {tk_rep:,.0f})</span></div>""", unsafe_allow_html=True)
        
        with st.expander("⚙️ ATUR REKAP (Sembunyikan Unit ke 'Lainnya')", expanded=False):
            st.write("Pilih Unit yang ingin digabung menjadi **'Lainnya'** di tabel Rekapitulasi.")
            if not df_keluar_all.empty:
                df_keluar_all['full_name'] = df_keluar_all['nama_alat'].astype(str) + " " + df_keluar_all['no_unit'].astype(str)
                unique_units = sorted(df_keluar_all['full_name'].unique().tolist())
                selected_excludes = st.multiselect("Pilih Unit:", unique_units, default=[u for u in unique_units if u in excluded_list])
                if st.button("Simpan Pengaturan Rekap"):
                    cursor.execute(f"DELETE FROM rekap_exclude WHERE lokasi_id={lokasi_id}")
                    for item in selected_excludes: cursor.execute("INSERT INTO rekap_exclude (lokasi_id, nama_unit_full) VALUES (%s, %s)", (lokasi_id, item))
                    conn.commit(); st.success("Pengaturan Disimpan!"); st.rerun()
            else: st.info("Belum ada data unit keluar.")
        
        with st.expander("🛠️ MENU ADMIN (Koreksi Nama/Unit)", expanded=False):
            t_kor, t_hap = st.tabs(["Koreksi Nama", "Hapus Data (Backup)"])
            with t_kor:
                c1, c2 = st.columns(2)
                with c1:
                    list_alat = sorted(df_keluar_all['nama_alat'].unique().tolist()) if not df_keluar_all.empty else []
                    if list_alat:
                        pilih_lama = st.selectbox("Alat Salah:", list_alat, key="ot"); input_baru = st.text_input("Nama Benar:", key="nt")
                        if st.button("Ganti Nama Alat"): 
                            cursor.execute(f"SELECT id FROM bbm_keluar WHERE nama_alat='{pilih_lama}' AND lokasi_id={lokasi_id}"); ids = [str(r[0]) for r in cursor.fetchall()]; ids_str = ",".join(ids)
                            if ids: cursor.execute("UPDATE bbm_keluar SET nama_alat=%s WHERE nama_alat=%s AND lokasi_id={lokasi_id}", (input_baru, pilih_lama, lokasi_id)); cursor.execute("INSERT INTO log_aktivitas (lokasi_id, kategori, deskripsi, affected_ids) VALUES (%s, %s, %s, %s)", (lokasi_id, "GANTI NAMA ALAT", f"Mengubah '{pilih_lama}' menjadi '{input_baru}'", ids_str)); conn.commit(); st.success("Nama Diganti & Dicatat!"); st.rerun()
                
                with c2:
                    list_alat_for_unit = sorted(df_keluar_all['nama_alat'].unique().tolist()) if not df_keluar_all.empty else []
                    if list_alat_for_unit:
                        pilih_alat_u = st.selectbox("Pilih Alat utk Ganti Unit:", list_alat_for_unit, key="pilih_alat_unit")
                        units_of_alat = sorted(df_keluar_all[df_keluar_all['nama_alat'] == pilih_alat_u]['no_unit'].unique().tolist())
                        pl_u = st.selectbox("Unit Salah:", units_of_alat, key="ou"); ib_u = st.text_input("Unit Benar:", key="nu")
                        if st.button("Ganti No Unit"):
                            cursor.execute(f"SELECT id FROM bbm_keluar WHERE no_unit='{pl_u}' AND nama_alat='{pilih_alat_u}' AND lokasi_id={lokasi_id}")
                            ids = [str(r[0]) for r in cursor.fetchall()]; ids_str = ",".join(ids)
                            if ids: 
                                cursor.execute("UPDATE bbm_keluar SET no_unit=%s WHERE no_unit=%s AND nama_alat=%s AND lokasi_id=%s", (ib_u, pl_u, pilih_alat_u, lokasi_id))
                                cursor.execute("INSERT INTO log_aktivitas (lokasi_id, kategori, deskripsi, affected_ids) VALUES (%s, %s, %s, %s)", (lokasi_id, "GANTI NO UNIT", f"Mengubah '{pl_u}' menjadi '{ib_u}' pada alat '{pilih_alat_u}'", ids_str)); conn.commit(); st.success("Unit Diganti & Dicatat!"); st.rerun()
                            else: st.warning("Data tidak ditemukan untuk kombinasi Alat dan Unit tersebut.")

            with t_hap:
                c1, c2 = st.columns(2)
                with c1:
                    if not df_masuk_all.empty:
                        m_sel = st.selectbox("Hapus Masuk:", df_masuk_all.apply(lambda x: f"{x['id']}|{x['tanggal']}|{x['sumber']}", axis=1))
                        if st.button("Hapus Masuk"): cursor.execute(f"DELETE FROM bbm_masuk WHERE id={m_sel.split('|')[0]}"); conn.commit(); st.rerun()
                with c2:
                    if not df_keluar_all.empty:
                        k_sel = st.selectbox("Hapus Keluar:", df_keluar_all.apply(lambda x: f"{x['id']}|{x['tanggal']}|{x['nama_alat']}", axis=1))
                        if st.button("Hapus Keluar"): cursor.execute(f"DELETE FROM bbm_keluar WHERE id={k_sel.split('|')[0]}"); conn.commit(); st.rerun()

        df_alat_g = pd.DataFrame(); df_truck_g = pd.DataFrame(); df_lain_g = pd.DataFrame()
        if not df_keluar_rep.empty:
            if 'kategori' not in df_keluar_rep.columns: df_keluar_rep['kategori'] = df_keluar_rep['nama_alat'].apply(cek_kategori)
            df_alat_g, df_truck_g, df_lain_g = segregate_data(df_keluar_rep, excluded_list)
            
            c_rekap1, c_rekap2, c_rekap3 = st.columns(3)
            with c_rekap1: 
                st.write("**Rekap Alat Berat**")
                if not df_alat_g.empty: st.dataframe(df_alat_g.groupby(['nama_alat', 'no_unit'])['jumlah_liter'].sum().reset_index(), hide_index=True)
            with c_rekap2: 
                st.write("**Rekap Mobil/Truck**")
                if not df_truck_g.empty: st.dataframe(df_truck_g.groupby(['nama_alat', 'no_unit'])['jumlah_liter'].sum().reset_index(), hide_index=True)
            with c_rekap3: 
                st.write("**Rekap Lainnya (Excluded)**")
                if not df_lain_g.empty: st.dataframe(df_lain_g.groupby(['nama_alat', 'no_unit'])['jumlah_liter'].sum().reset_index(), hide_index=True)

        st.divider(); col_a, col_b = st.columns(2)
        with col_a: st.subheader("📋 Data Masuk BBM (Periode Ini)"); st.dataframe(df_masuk_rep.sort_values('tanggal', ascending=False), use_container_width=True)
        with col_b: st.subheader("📋 Data Penggunaan BBM (Periode Ini)"); st.dataframe(df_keluar_rep.sort_values('tanggal', ascending=False), use_container_width=True)
        st.divider(); st.subheader("📊 Monitoring (Periode Ini)"); 
        img_buffer = generate_chart_for_report(df_alat_g, df_truck_g)
        if img_buffer: st.image(img_buffer, caption="Diagram Alat Berat & Mobil")
        else: st.info("Belum ada data untuk ditampilkan di grafik.")

        st.divider(); st.subheader(f"📅 Monitoring Bulanan ({start_rep.strftime('%b %Y')} - {end_rep.strftime('%b %Y')})")
        stok_run_m = hitung_stok_awal_periode(conn, lokasi_id, start_rep)
        m_data_mon = []
        curr = start_rep.replace(day=1); end_limit = end_rep.replace(day=1)
        while curr <= end_limit:
            m = curr.month; y = curr.year
            q_in = f"SELECT SUM(jumlah_liter) FROM bbm_masuk WHERE lokasi_id={lokasi_id} AND MONTH(tanggal)={m} AND YEAR(tanggal)={y}"
            q_out = f"SELECT SUM(jumlah_liter) FROM bbm_keluar WHERE lokasi_id={lokasi_id} AND MONTH(tanggal)={m} AND YEAR(tanggal)={y}"
            cursor = conn.cursor()
            cursor.execute(q_in); res_in = cursor.fetchone(); mi = float(res_in[0]) if res_in and res_in[0] else 0.0
            cursor.execute(q_out); res_out = cursor.fetchone(); mo = float(res_out[0]) if res_out and res_out[0] else 0.0
            prev = stok_run_m; stok_run_m = prev + mi - mo
            m_data_mon.append({'Bulan': curr.strftime("%B %Y"), 'Sisa Bulan Lalu': prev, 'Masuk': mi, 'Keluar': mo, 'Sisa Akhir': stok_run_m, 'bulan_nama': curr.strftime("%b")})
            curr = curr + relativedelta(months=1)

        df_m_mon = pd.DataFrame(m_data_mon)
        c_g1, c_g2 = st.columns(2)
        with c_g1: 
            if not df_m_mon.empty:
                df_for_graph = df_m_mon.rename(columns={'Masuk': 'masuk', 'Keluar': 'keluar'})
                img_m = generate_monthly_chart(df_for_graph)
                if img_m: st.image(img_m)
            else: st.info("Tidak ada data dalam range ini.")
        with c_g2: st.dataframe(df_m_mon[['Bulan', 'Sisa Bulan Lalu', 'Masuk', 'Keluar', 'Sisa Akhir']], hide_index=True)

    with t3:
        st.header("🖨️ Export Laporan Periode")
        st.write("Silakan pilih periode laporan. Sistem akan membuat laporan **Bulan demi Bulan** secara otomatis.")
        
        # Pilihan Mode Export
        export_mode = st.radio("Pilih Mode Export:", ["📄 Standard (A4 Berhalaman)", "📜 1 Bulan 1 Kertas (Custom)"], horizontal=True)
        
        if 't3_start' not in st.session_state: st.session_state.t3_start = first_day
        if 't3_end' not in st.session_state: st.session_state.t3_end = today
        c_d1, c_d2 = st.columns(2)
        with c_d1: start_date_exp = st.date_input("Dari Tanggal", value=st.session_state.t3_start, key="pick_start_t3"); st.session_state.t3_start = start_date_exp
        with c_d2: end_date_exp = st.date_input("Sampai Tanggal", value=st.session_state.t3_end, key="pick_end_t3"); st.session_state.t3_end = end_date_exp

        c1, c2, c3 = st.columns(3)
        if start_date_exp <= end_date_exp:
            with c1: 
                if st.button("📕 Download PDF", use_container_width=True): 
                    if "1 Bulan 1 Kertas" in export_mode:
                        pdf = generate_pdf_one_sheet(conn, lokasi_id, nama_proyek, start_date_exp, end_date_exp, excluded_list)
                    else:
                        pdf = generate_pdf_portrait(conn, lokasi_id, nama_proyek, start_date_exp, end_date_exp, excluded_list)
                    st.download_button("⬇️ Simpan PDF", pdf, f"Laporan_{nama_proyek}_{start_date_exp}_{end_date_exp}.pdf", "application/pdf")
            with c2: 
                if st.button("📗 Download Excel", use_container_width=True): 
                    if "1 Bulan 1 Kertas" in export_mode:
                        xl = generate_excel_one_sheet(conn, lokasi_id, nama_proyek, start_date_exp, end_date_exp, excluded_list)
                    else:
                        xl = generate_excel_styled(conn, lokasi_id, nama_proyek, start_date_exp, end_date_exp, excluded_list)
                    st.download_button("⬇️ Simpan Excel", xl, f"Laporan_{nama_proyek}_{start_date_exp}_{end_date_exp}.xlsx")
            with c3: 
                if st.button("📘 Download Word", use_container_width=True): 
                    if "1 Bulan 1 Kertas" in export_mode:
                        doc = generate_docx_one_sheet(conn, lokasi_id, nama_proyek, start_date_exp, end_date_exp, excluded_list)
                    else:
                        doc = generate_docx_fixed(conn, lokasi_id, nama_proyek, start_date_exp, end_date_exp, excluded_list)
                    st.download_button("⬇️ Simpan Word", doc, f"Laporan_{nama_proyek}_{start_date_exp}_{end_date_exp}.docx")
        else: st.error("Tanggal Akhir harus lebih besar dari Tanggal Awal")

if __name__ == "__main__":
    main()